"""
通用表單結構分析器 — Excel / Word

掃描檔案，識別出所有可能的欄位標籤，並找出對應的值位置（空白儲存格或佔位符）。
輸出 field_map 作為後續 AI 映射與自動回填的基礎。

不含任何 domain 知識；欄位關鍵字可由呼叫端客製。
"""

import io
import logging
from typing import Iterable, Optional

from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from docx import Document

from app.autofill_core.field_detection import (
    is_field_label,
    is_placeholder,
    guess_field_type,
)

logger = logging.getLogger(__name__)


class StructureAnalyzer:
    """Excel/Word 表單結構分析器"""

    def __init__(
        self,
        field_keywords: Optional[Iterable[str]] = None,
        excel_max_rows: int = 200,
        excel_max_cols: int = 50,
    ):
        """
        Args:
            field_keywords: 自訂欄位標籤關鍵字；None 則使用內建通用集。
            excel_max_rows / excel_max_cols: 掃描範圍上限。
        """
        self._keywords = tuple(field_keywords) if field_keywords is not None else None
        self._max_rows = excel_max_rows
        self._max_cols = excel_max_cols

    # ================================================================
    # Public API
    # ================================================================

    async def analyze(self, file_content: bytes, file_name: str) -> list[dict]:
        """分析檔案並回傳 field_map。"""
        file_type = file_name.split('.')[-1].lower()
        if file_type == 'xlsx':
            return await self.analyze_excel(file_content)
        if file_type == 'docx':
            return await self.analyze_word(file_content)
        raise ValueError(f"不支援的檔案格式: {file_type}")

    async def extract_text(self, file_content: bytes, file_name: str) -> str:
        """擷取檔案文字內容（供 AI 上下文用）。"""
        file_type = file_name.split('.')[-1].lower()
        if file_type == 'xlsx':
            return await self._extract_excel_text(file_content)
        if file_type == 'docx':
            return await self._extract_word_text(file_content)
        raise ValueError(f"不支援的檔案格式: {file_type}")

    # ================================================================
    # Excel
    # ================================================================

    async def analyze_excel(self, content: bytes) -> list[dict]:
        fields: list[dict] = []
        wb = load_workbook(io.BytesIO(content))

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            merged_ranges = list(ws.merged_cells.ranges)
            merge_lookup: dict = {}
            for mr in merged_ranges:
                for row in range(mr.min_row, mr.max_row + 1):
                    for col in range(mr.min_col, mr.max_col + 1):
                        coord = f"{get_column_letter(col)}{row}"
                        merge_lookup[coord] = {
                            "range": str(mr),
                            "top_left": f"{get_column_letter(mr.min_col)}{mr.min_row}",
                            "rows": mr.max_row - mr.min_row + 1,
                            "cols": mr.max_col - mr.min_col + 1,
                        }

            max_row = min(ws.max_row or 1, self._max_rows)
            max_col = min(ws.max_column or 1, self._max_cols)

            for row_idx in range(1, max_row + 1):
                for col_idx in range(1, max_col + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if cell.value is None:
                        continue

                    value = str(cell.value).strip()
                    if not value:
                        continue

                    coord = f"{get_column_letter(col_idx)}{row_idx}"

                    if coord in merge_lookup and merge_lookup[coord]["top_left"] != coord:
                        continue

                    if not is_field_label(value, keywords=self._keywords):
                        continue

                    value_cell = self._find_value_cell_excel(
                        ws, row_idx, col_idx, max_row, max_col, merge_lookup
                    )

                    fields.append({
                        "field_id": f"excel_{sheet_name}_{coord}",
                        "field_name": value.rstrip(':：ˍ_ '),
                        "field_type": guess_field_type(value),
                        "label_location": {
                            "sheet": sheet_name,
                            "cell": coord,
                            "row": row_idx,
                            "column": col_idx,
                        },
                        "value_location": value_cell,
                        "is_merged": coord in merge_lookup,
                        "merge_info": merge_lookup.get(coord),
                        "mapping": None,
                    })

        return fields

    def _find_value_cell_excel(
        self,
        ws,
        label_row: int,
        label_col: int,
        max_row: int,
        max_col: int,
        merge_lookup: dict,
    ) -> Optional[dict]:
        """尋找標籤對應的值儲存格（先右再下）。"""
        sheet_name = ws.title
        label_coord = f"{get_column_letter(label_col)}{label_row}"

        label_merge = merge_lookup.get(label_coord)
        if label_merge:
            merge_max_col = label_col + label_merge["cols"] - 1
            merge_max_row = label_row + label_merge["rows"] - 1
        else:
            merge_max_col = label_col
            merge_max_row = label_row

        # 右側搜尋
        search_start_col = merge_max_col + 1
        for next_col in range(search_start_col, min(search_start_col + 3, max_col + 1)):
            next_coord = f"{get_column_letter(next_col)}{label_row}"
            next_cell = ws.cell(row=label_row, column=next_col)

            target_merge = merge_lookup.get(next_coord)
            actual_coord = target_merge["top_left"] if target_merge else next_coord

            cell_val = next_cell.value
            if cell_val is None or is_placeholder(str(cell_val)):
                return {
                    "sheet": sheet_name,
                    "cell": actual_coord,
                    "row": label_row,
                    "column": next_col,
                    "direction": "right",
                    "offset": next_col - label_col,
                }

            if not is_field_label(str(cell_val), keywords=self._keywords):
                return {
                    "sheet": sheet_name,
                    "cell": actual_coord,
                    "row": label_row,
                    "column": next_col,
                    "direction": "right",
                    "offset": next_col - label_col,
                }

        # 下方搜尋
        search_start_row = merge_max_row + 1
        for next_row in range(search_start_row, min(search_start_row + 2, max_row + 1)):
            next_coord = f"{get_column_letter(label_col)}{next_row}"
            next_cell = ws.cell(row=next_row, column=label_col)

            target_merge = merge_lookup.get(next_coord)
            actual_coord = target_merge["top_left"] if target_merge else next_coord

            cell_val = next_cell.value
            if cell_val is None or is_placeholder(str(cell_val)):
                return {
                    "sheet": sheet_name,
                    "cell": actual_coord,
                    "row": next_row,
                    "column": label_col,
                    "direction": "below",
                    "offset": next_row - label_row,
                }

        return None

    async def _extract_excel_text(self, content: bytes) -> str:
        lines: list[str] = []
        wb = load_workbook(io.BytesIO(content))
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            max_row = min(ws.max_row or 1, 100)
            max_col = min(ws.max_column or 1, 30)
            for row_idx in range(1, max_row + 1):
                row_texts: list[str] = []
                for col_idx in range(1, max_col + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if cell.value is not None:
                        row_texts.append(str(cell.value).strip())
                if row_texts:
                    lines.append(" | ".join(row_texts))
        return "\n".join(lines)

    # ================================================================
    # Word
    # ================================================================

    async def analyze_word(self, content: bytes) -> list[dict]:
        fields: list[dict] = []
        doc = Document(io.BytesIO(content))

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            if is_field_label(text, keywords=self._keywords) or '____' in text or '＿＿' in text:
                field_name = text.split(':')[0].split('：')[0].strip().rstrip('_＿ ')

                fields.append({
                    "field_id": f"word_para_{para_idx}",
                    "field_name": field_name,
                    "field_type": guess_field_type(field_name),
                    "label_location": {
                        "type": "paragraph",
                        "paragraph_index": para_idx,
                    },
                    "value_location": {
                        "type": "paragraph",
                        "paragraph_index": para_idx,
                        "replace_pattern": "after_colon",
                    },
                    "mapping": None,
                })

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text or not is_field_label(text, keywords=self._keywords):
                        continue

                    field_name = text.rstrip(':：_＿ ')
                    value_loc = self._find_value_cell_word(table, row_idx, cell_idx, table_idx)

                    fields.append({
                        "field_id": f"word_t{table_idx}_r{row_idx}_c{cell_idx}",
                        "field_name": field_name,
                        "field_type": guess_field_type(field_name),
                        "label_location": {
                            "type": "table",
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "cell_index": cell_idx,
                        },
                        "value_location": value_loc,
                        "mapping": None,
                    })

        return fields

    def _find_value_cell_word(
        self, table, label_row: int, label_col: int, table_idx: int = 0
    ) -> Optional[dict]:
        rows = table.rows
        num_cols = len(rows[label_row].cells) if label_row < len(rows) else 0

        if label_col + 1 < num_cols:
            right_cell = rows[label_row].cells[label_col + 1]
            right_text = right_cell.text.strip()
            if not right_text or is_placeholder(right_text):
                return {
                    "type": "table",
                    "table_index": table_idx,
                    "row_index": label_row,
                    "cell_index": label_col + 1,
                    "direction": "right",
                }

        if label_row + 1 < len(rows):
            below_cell = rows[label_row + 1].cells[label_col]
            below_text = below_cell.text.strip()
            if not below_text or is_placeholder(below_text):
                return {
                    "type": "table",
                    "table_index": table_idx,
                    "row_index": label_row + 1,
                    "cell_index": label_col,
                    "direction": "below",
                }

        return None

    async def _extract_word_text(self, content: bytes) -> str:
        lines: list[str] = []
        doc = Document(io.BytesIO(content))
        for para in doc.paragraphs[:100]:
            text = para.text.strip()
            if text:
                lines.append(text)
        for table_idx, table in enumerate(doc.tables):
            lines.append(f"=== 表格 {table_idx + 1} ===")
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    lines.append(" | ".join(row_texts))
        return "\n".join(lines)
