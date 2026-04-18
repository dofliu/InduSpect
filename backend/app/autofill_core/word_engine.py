"""
Word 自動回填引擎 — 通用 python-docx 操作

負責將 fill_values 寫入 docx 檔案的段落或表格儲存格位置，保留格式。
不含任何 domain 邏輯。
"""

import io
import logging

from docx import Document

from app.autofill_core.field_detection import replace_paragraph_text_preserve_format

logger = logging.getLogger(__name__)


class WordAutoFillEngine:
    """Word 表單自動回填引擎"""

    async def fill(
        self,
        file_content: bytes,
        field_lookup: dict,
        value_lookup: dict,
    ) -> bytes:
        """將 value_lookup 中的值寫入 field_lookup 指定的段落/表格位置。"""
        doc = Document(io.BytesIO(file_content))

        for field_id, value in value_lookup.items():
            field = field_lookup.get(field_id)
            if not field:
                continue

            val_loc = field.get("value_location")
            if not val_loc:
                continue

            loc_type = val_loc.get("type")

            if loc_type == "paragraph":
                self._fill_paragraph(doc, val_loc, value)
            elif loc_type == "table":
                self._fill_table_cell(doc, field, val_loc, value)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _fill_paragraph(self, doc, val_loc: dict, value) -> None:
        para_idx = val_loc.get("paragraph_index")
        if para_idx is None or para_idx >= len(doc.paragraphs):
            return

        para = doc.paragraphs[para_idx]
        replace_pattern = val_loc.get("replace_pattern", "after_colon")

        if replace_pattern == "after_colon":
            text = para.text
            for sep in ['：', ':']:
                if sep in text:
                    prefix = text.split(sep)[0] + sep
                    replace_paragraph_text_preserve_format(
                        para, f"{prefix} {value}"
                    )
                    break
        else:
            replace_paragraph_text_preserve_format(para, value)

    def _fill_table_cell(self, doc, field: dict, val_loc: dict, value) -> None:
        table_idx = val_loc.get("table_index")
        row_idx = val_loc.get("row_index")
        cell_idx = val_loc.get("cell_index")

        # 向後相容：舊版 field_map 的 table_index 可能為 None，
        # 用 label_location 的 table_index 作為 fallback
        if table_idx is None:
            label_loc = field.get("label_location", {})
            table_idx = label_loc.get("table_index")

        if not (
            table_idx is not None
            and table_idx < len(doc.tables)
            and row_idx is not None
            and cell_idx is not None
        ):
            return

        table = doc.tables[table_idx]
        if row_idx >= len(table.rows):
            return
        row = table.rows[row_idx]
        if cell_idx >= len(row.cells):
            return

        cell = row.cells[cell_idx]
        if cell.paragraphs:
            replace_paragraph_text_preserve_format(
                cell.paragraphs[0], str(value)
            )
        else:
            cell.text = str(value)
