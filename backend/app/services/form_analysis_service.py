"""
表單分析服務 — 結構分析、模板建立、欄位映射

職責：
- 深度分析 Excel/Word 表格結構（委派給 autofill_core.StructureAnalyzer）
- 從真實表單自動建立 InspectionTemplate（工業巡檢專屬）
- AI 智慧欄位映射（工業巡檢專屬）
- 精準欄位映射（基於 photo_task_bindings，工業巡檢專屬）
"""

import io
import re
import json
import uuid
import logging
from typing import Optional
from datetime import datetime

import google.generativeai as genai
from openpyxl import load_workbook
from docx import Document

from app.config import settings
from app.constants import FIELD_KEYWORDS
from app.autofill_core import StructureAnalyzer
from app.services.form_utils import is_non_field_item, guess_field_type

logger = logging.getLogger(__name__)


class FormAnalysisService:
    """表單分析與模板建立服務（InduSpect 工業巡檢專用）"""

    def __init__(self):
        genai.configure(api_key=settings.gemini_api_key)
        # 通用結構分析器（使用工業巡檢關鍵字擴充預設集）
        self._analyzer = StructureAnalyzer(field_keywords=FIELD_KEYWORDS)

    # ================================================================
    # 精準欄位映射
    # ================================================================

    async def precision_map_fields(
        self,
        field_map: list[dict],
        inspection_results: list[dict],
        photo_task_bindings: list[dict],
    ) -> dict:
        """
        利用 photo_task_bindings 進行精準映射

        每個 binding 已知「哪張照片對應哪些欄位」，
        因此不需要 AI 猜測映射關係，直接對應。

        對於 binding 未覆蓋的欄位（基本資訊等），
        退回通用 AI 映射或使用預設值。
        """
        mappings = []
        covered_field_ids = set()

        # Phase 1: 從 photo_task_bindings 精準映射
        for binding in photo_task_bindings:
            ai_result = binding.get("ai_result")
            if not ai_result:
                continue

            task_id = binding.get("task_id", "")
            value_field_ids = binding.get("value_field_ids", [])
            judgment_field_ids = binding.get("judgment_field_ids", [])
            remarks_field_ids = binding.get("remarks_field_ids", [])

            # 從 ai_result 取得分析數據
            readings = ai_result.get("readings", [])
            is_anomaly = ai_result.get("is_anomaly", False)
            condition = ai_result.get("condition_assessment", "")
            anomaly_desc = ai_result.get("anomaly_description", "")
            summary = ai_result.get("summary", "")

            # 1. 數值欄位 → 填入 readings
            for idx, fid in enumerate(value_field_ids):
                field_info = self._find_field_in_map(field_map, fid)
                if not field_info:
                    continue
                fname = field_info.get("field_name", "")
                ftype = field_info.get("field_type", "text")

                if ftype == "number" and readings:
                    # 嘗試匹配最相關的 reading
                    matched = self._match_reading(fname, readings)
                    if matched:
                        mappings.append({
                            "field_id": fid,
                            "suggested_value": str(matched.get("value", "")),
                            "source": f"[{task_id}] AI 讀取: {matched.get('label', '')} = {matched.get('value', '')}{matched.get('unit', '')}",
                            "confidence": 0.95,
                        })
                        covered_field_ids.add(fid)
                        continue

                # 非數值 → 填入狀態描述
                if condition:
                    mappings.append({
                        "field_id": fid,
                        "suggested_value": condition,
                        "source": f"[{task_id}] AI 狀態評估",
                        "confidence": 0.85,
                    })
                    covered_field_ids.add(fid)

            # 2. 判定欄位 → 填入合格/不合格
            for fid in judgment_field_ids:
                judgment_value = "不合格" if is_anomaly else "合格"
                mappings.append({
                    "field_id": fid,
                    "suggested_value": judgment_value,
                    "source": f"[{task_id}] AI 判定: is_anomaly={is_anomaly}",
                    "confidence": 0.92,
                })
                covered_field_ids.add(fid)

            # 3. 備註欄位 → 異常時填入異常描述
            for fid in remarks_field_ids:
                if is_anomaly and anomaly_desc:
                    mappings.append({
                        "field_id": fid,
                        "suggested_value": anomaly_desc,
                        "source": f"[{task_id}] AI 異常描述",
                        "confidence": 0.88,
                    })
                else:
                    mappings.append({
                        "field_id": fid,
                        "suggested_value": summary if summary else "正常",
                        "source": f"[{task_id}] AI 摘要",
                        "confidence": 0.80,
                    })
                covered_field_ids.add(fid)

        # Phase 2: 未覆蓋的欄位 → 使用通用資訊或 AI 映射
        uncovered_fields = [
            f for f in field_map if f["field_id"] not in covered_field_ids
        ]

        if uncovered_fields and inspection_results:
            # 嘗試用通用資訊填入（日期、人員等）
            general_info = inspection_results[0] if inspection_results else {}
            for field in uncovered_fields:
                fid = field["field_id"]
                fname = field.get("field_name", "")

                # 嘗試從 general_info 直接匹配
                matched_value = self._match_general_info(fname, general_info)
                if matched_value is not None:
                    mappings.append({
                        "field_id": fid,
                        "suggested_value": str(matched_value),
                        "source": "基本資訊直接匹配",
                        "confidence": 0.90,
                    })
                    covered_field_ids.add(fid)

        return {
            "success": True,
            "mappings": mappings,
            "unmapped_fields": [
                f["field_id"] for f in field_map
                if f["field_id"] not in covered_field_ids
            ],
        }

    def _find_field_in_map(self, field_map: list[dict], field_id: str) -> Optional[dict]:
        """在 field_map 中查找指定 field_id 的欄位"""
        for f in field_map:
            if f["field_id"] == field_id:
                return f
        return None

    def _match_reading(self, field_name: str, readings: list[dict]) -> Optional[dict]:
        """從 AI readings 中匹配最相關的讀數"""
        if not readings:
            return None

        # 精確匹配
        for r in readings:
            label = r.get("label", "")
            if label and label in field_name:
                return r
            if field_name and field_name in label:
                return r

        # 單位匹配
        unit_keywords = {
            'V': ['電壓'], 'A': ['電流'], 'Ω': ['電阻', '接地'],
            'MΩ': ['絕緣'], '°C': ['溫度'], 'Hz': ['頻率'],
            'rpm': ['轉速'], 'mm/s': ['振動'], 'dB': ['噪音'],
        }
        for r in readings:
            r_unit = r.get("unit", "")
            if r_unit in unit_keywords:
                for kw in unit_keywords[r_unit]:
                    if kw in field_name:
                        return r

        # 只有一個 reading → 直接使用
        if len(readings) == 1:
            return readings[0]

        return None

    def _match_general_info(self, field_name: str, info: dict) -> Optional[str]:
        """從通用檢查資訊中匹配基本欄位值"""
        match_rules = {
            '設備名稱': 'equipment_name',
            '設備類型': 'equipment_type',
            '設備編號': 'equipment_id',
            '檢查日期': 'inspection_date',
            '日期': 'inspection_date',
            '檢查人員': 'inspector_name',
            '人員': 'inspector_name',
            '姓名': 'inspector_name',
            '位置': 'location',
            '地點': 'location',
            '廠區': 'location',
        }

        for keyword, info_key in match_rules.items():
            if keyword in field_name:
                value = info.get(info_key)
                if value is not None and str(value).strip():
                    return str(value)

        return None

    # ================================================================
    # 動態模板建立（從真實表單自動產生 InspectionTemplate）
    # ================================================================

    async def create_template_from_file(
        self,
        file_content: bytes,
        file_name: str,
        template_name: str,
        category: str = "一般設備",
        company: str = "",
        department: str = "",
    ) -> dict:
        """
        從真實廠商 Excel/Word 表單自動建立 InspectionTemplate JSON

        流程：
        1. 深度分析表格結構 → 取得 field_map
        2. 用 Gemini AI 將 field_map 轉換為 InspectionTemplate 格式
        3. 儲存原始文件 + 產生的模板

        Returns:
            包含完整 InspectionTemplate JSON 的 dict
        """
        file_type = file_name.split('.')[-1].lower()

        # Step 1: 深度分析表格結構（委派給 autofill_core）
        field_map = await self._analyzer.analyze(file_content, file_name)
        raw_text = await self._analyzer.extract_text(file_content, file_name)

        if not field_map:
            raise ValueError("無法從檔案中識別出任何欄位，請確認檔案格式正確")

        # Step 2: 用 AI 轉換為 InspectionTemplate 格式
        template_json = await self._ai_convert_to_inspection_template(
            field_map=field_map,
            raw_text=raw_text,
            template_name=template_name,
            category=category,
            company=company,
            department=department,
        )

        # Step 3: 綁定原始文件資訊（供日後回填使用）
        template_id = template_json.get("template_id", f"TEMP-{uuid.uuid4().hex[:8].upper()}")
        template_json["template_id"] = template_id
        template_json["source_file"] = {
            "file_name": file_name,
            "file_type": file_type,
            "field_map": field_map,
        }

        return {
            "success": True,
            "template_id": template_id,
            "template": template_json,
            "field_count": sum(
                len(s.get("fields", []))
                for s in template_json.get("sections", [])
            ),
            "section_count": len(template_json.get("sections", [])),
            "message": f"成功從 {file_name} 建立模板「{template_name}」",
            # Also return field_map for storage
            "field_map": field_map,
            "file_type": file_type,
        }

    async def _ai_convert_to_inspection_template(
        self,
        field_map: list[dict],
        raw_text: str,
        template_name: str,
        category: str,
        company: str,
        department: str,
    ) -> dict:
        """
        使用 Gemini AI 將分析出的欄位地圖轉換為完整的 InspectionTemplate JSON
        """
        # 準備欄位摘要
        field_summary = []
        for f in field_map:
            field_summary.append({
                "field_id": f["field_id"],
                "field_name": f["field_name"],
                "guessed_type": f.get("field_type", "text"),
            })

        try:
            model = genai.GenerativeModel(settings.gemini_flash_model)

            prompt = f"""你是一位工業定檢表單分析專家。請根據以下從真實定檢表格中擷取出的欄位資訊，
建立一個完整的 InspectionTemplate JSON。

【表單名稱】{template_name}
【設備類別】{category}
【公司名稱】{company}
【部門】{department}

【偵測到的欄位列表】
{json.dumps(field_summary, ensure_ascii=False, indent=2)}

【表格原始文字內容（前 2000 字）】
{raw_text[:2000]}

請產生一個完整的 InspectionTemplate JSON，格式要求如下：

{{
  "template_id": "TEMP-自動生成ID",
  "template_name": "{template_name}",
  "template_version": "1.0",
  "category": "{category}",
  "created_at": "{datetime.now().isoformat()}",
  "updated_at": "{datetime.now().isoformat()}",
  "metadata": {{
    "company": "{company}",
    "department": "{department}",
    "inspection_cycle_days": 30,
    "estimated_duration_minutes": 依表單複雜度估算,
    "required_tools": ["相機"],
    "safety_notes": "依表單內容判斷"
  }},
  "sections": [
    {{
      "section_id": "英文ID",
      "section_title": "中文標題",
      "section_order": 1,
      "description": "簡短說明",
      "fields": [
        {{
          "field_id": "使用原始 field_id 或語意化 ID",
          "field_type": "text|number|radio|checkbox|dropdown|datetime|date|photo|textarea|signature",
          "label": "欄位標籤",
          "placeholder": "提示文字（可選）",
          "required": true/false,
          "ai_fillable": true/false,
          "unit": "單位（可選，數值欄位用）",
          "options": [只有 radio/checkbox/dropdown 才需要],
          "validation": {{可選}},
          "warning_threshold": {{可選，數值欄位用}}
        }}
      ]
    }}
  ]
}}

【重要規則】
1. 將欄位合理分組為 sections：
   - 基本資訊（設備名稱、編號、日期、人員等）
   - 各類檢測項目（按邏輯分組）
   - 綜合評估/結論
2. field_type 判斷規則：
   - 包含「日期」→ datetime 或 date
   - 包含「溫度/壓力/電流/電壓/讀數」等量測值 → number（設定 unit 和 validation）
   - 包含「是否/合格/正常/異常/判定」→ radio（提供 options）
   - 包含「狀態/狀況/等級」→ radio 或 dropdown
   - 包含「備註/說明/描述」→ textarea
   - 包含「姓名/簽名」且在表尾 → signature
   - 其他 → text
3. 量測值欄位標記 ai_fillable: true
4. 每個 section 至少包含一個照片欄位（field_type: "photo"），用於 AI 拍照辨識
5. 最後的 section 應包含一個 signature 欄位
6. options 格式: [{{"value": "英文值", "label": "中文標籤"}}]

只回應 JSON，不要其他文字。"""

            response = model.generate_content(prompt)

            # 解析 JSON
            json_match = re.search(r'\{[\s\S]*\}', response.text)
            if json_match:
                template_json = json.loads(json_match.group())
                if template_json.get('sections') and len(template_json['sections']) > 0:
                    return template_json
                else:
                    logger.warning("AI 回傳的模板 JSON 缺少 sections，使用 fallback")

        except json.JSONDecodeError as e:
            logger.error(f"AI 回應 JSON 解析失敗: {e}")
        except Exception as e:
            logger.error(f"AI 轉換模板失敗: {e}")

        # Fallback: 如果 AI 失敗，用規則產生基本模板
        return self._fallback_create_template(
            field_map, template_name, category, company, department
        )

    def _fallback_create_template(
        self,
        field_map: list[dict],
        template_name: str,
        category: str,
        company: str,
        department: str,
    ) -> dict:
        """
        AI 失敗時的 fallback：用規則將 field_map 轉換為基本 InspectionTemplate
        """
        now = datetime.now().isoformat()
        template_id = f"TEMP-{uuid.uuid4().hex[:8].upper()}"

        # 分類欄位
        basic_fields = []
        inspection_fields = []
        measurement_fields = []
        conclusion_fields = []

        basic_kw = ['名稱', '編號', '日期', '人員', '姓名', '位置', '地點',
                     '廠商', '製造', '型號', '規格', '電話', '證照', '週期',
                     '樓層', '區域', '天氣', '時間', '陪同']
        measure_kw = ['電壓', '電流', '溫度', '壓力', '電阻', '絕緣', '頻率',
                       '振動', '噪音', '轉速', '流量', '油位', '水位', '濕度',
                       '功率', '相位', 'R相', 'S相', 'T相']
        conclusion_kw = ['綜合', '結論', '改善', '評估', '簽名', '簽核',
                          '複查', '整體', '判定結果']

        for f in field_map:
            name = f["field_name"]

            if is_non_field_item(name):
                continue

            field_type = f.get("field_type", "text")
            field_entry = {
                "field_id": f["field_id"],
                "field_type": field_type,
                "label": name,
                "required": False,
                "ai_fillable": field_type in ("number", "checkbox"),
            }

            if any(kw in name for kw in ['合格', '判定', '正常', '異常']):
                field_entry["field_type"] = "radio"
                field_entry["options"] = [
                    {"value": "normal", "label": "正常"},
                    {"value": "abnormal", "label": "異常"},
                    {"value": "na", "label": "不適用"},
                ]

            if field_type == "number" or any(kw in name for kw in measure_kw):
                field_entry["field_type"] = "number"
                field_entry["ai_fillable"] = True
                if '電壓' in name:
                    field_entry["unit"] = "V"
                elif '電流' in name:
                    field_entry["unit"] = "A"
                elif '電阻' in name:
                    field_entry["unit"] = "Ω"
                elif '絕緣' in name:
                    field_entry["unit"] = "MΩ"
                elif '溫度' in name:
                    field_entry["unit"] = "°C"
                elif '壓力' in name:
                    field_entry["unit"] = "kPa"

            if field_type == "date":
                field_entry["field_type"] = "date"
                field_entry["ai_fillable"] = True

            if any(kw in name for kw in ['備註', '說明', '描述', '建議']):
                field_entry["field_type"] = "textarea"
                field_entry["rows"] = 3

            if any(kw in name for kw in basic_kw):
                basic_fields.append(field_entry)
            elif any(kw in name for kw in measure_kw):
                measurement_fields.append(field_entry)
            elif any(kw in name for kw in conclusion_kw):
                conclusion_fields.append(field_entry)
            else:
                inspection_fields.append(field_entry)

        # 構建 sections
        sections = []
        order = 1

        if basic_fields:
            sections.append({
                "section_id": "basic_info",
                "section_title": "基本資訊",
                "section_order": order,
                "description": "設備基本資料與檢測資訊",
                "fields": basic_fields,
            })
            order += 1

        if inspection_fields:
            inspection_fields.append({
                "field_id": "inspection_photo",
                "field_type": "photo",
                "label": "檢測現場照片",
                "required": False,
                "ai_analyze": True,
            })
            sections.append({
                "section_id": "inspection_items",
                "section_title": "檢測項目",
                "section_order": order,
                "description": "設備外觀與功能檢測",
                "fields": inspection_fields,
            })
            order += 1

        if measurement_fields:
            measurement_fields.append({
                "field_id": "measurement_photo",
                "field_type": "photo",
                "label": "量測儀表照片",
                "required": False,
                "ai_analyze": True,
            })
            sections.append({
                "section_id": "measurements",
                "section_title": "量測數據",
                "section_order": order,
                "description": "各項量測數值與讀數",
                "fields": measurement_fields,
            })
            order += 1

        conclusion_section_fields = list(conclusion_fields)
        has_overall = any('判定' in f.get('label', '') or '評估' in f.get('label', '')
                          for f in conclusion_section_fields)
        if not has_overall:
            conclusion_section_fields.insert(0, {
                "field_id": "overall_result",
                "field_type": "radio",
                "label": "整體判定",
                "required": True,
                "ai_fillable": False,
                "options": [
                    {"value": "pass", "label": "合格"},
                    {"value": "fail", "label": "不合格"},
                    {"value": "conditional", "label": "有條件通過"},
                ],
            })
        has_notes = any('備註' in f.get('label', '') for f in conclusion_section_fields)
        if not has_notes:
            conclusion_section_fields.append({
                "field_id": "notes",
                "field_type": "textarea",
                "label": "備註說明",
                "required": False,
                "ai_fillable": False,
                "max_length": 500,
                "rows": 4,
            })
        conclusion_section_fields.append({
            "field_id": "inspector_signature",
            "field_type": "signature",
            "label": "檢測人員簽名",
            "required": True,
            "save_as_image": True,
        })

        sections.append({
            "section_id": "conclusion",
            "section_title": "綜合評估",
            "section_order": order,
            "description": "檢測結論與簽名",
            "fields": conclusion_section_fields,
        })

        total_fields = sum(len(s["fields"]) for s in sections)

        return {
            "template_id": template_id,
            "template_name": template_name,
            "template_version": "1.0",
            "category": category,
            "created_at": now,
            "updated_at": now,
            "metadata": {
                "company": company,
                "department": department,
                "inspection_cycle_days": 30,
                "estimated_duration_minutes": max(15, total_fields * 2),
                "required_tools": ["相機"],
            },
            "sections": sections,
        }

    # ================================================================
    # 模板分析
    # ================================================================

    async def analyze_template(
        self,
        file_content: bytes,
        file_name: str,
        vendor_name: str,
        template_name: str,
        description: Optional[str] = None
    ) -> dict:
        """
        使用 AI 分析模板結構

        自動識別欄位並建議與巡檢資料的對應關係
        """
        template_id = str(uuid.uuid4())
        file_type = file_name.split('.')[-1].lower()

        if file_type == 'xlsx':
            fields, raw_structure = await self._parse_excel_template(file_content)
        elif file_type == 'docx':
            fields, raw_structure = await self._parse_word_template(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        suggested_mappings = await self._ai_suggest_mappings(fields, raw_structure)

        template = {
            "id": template_id,
            "name": template_name,
            "vendor_name": vendor_name,
            "file_type": file_type,
            "description": description,
            "fields": fields,
            "file_content": file_content,
            "created_at": datetime.now().isoformat(),
        }

        return {
            "success": True,
            "template_id": template_id,
            "detected_fields": fields,
            "suggested_mappings": suggested_mappings,
            "message": f"成功分析模板，識別到 {len(fields)} 個欄位",
            "_template": template,  # For storage by orchestrator
        }

    # ================================================================
    # 深度結構分析 — 委派給 autofill_core.StructureAnalyzer
    # ================================================================

    async def analyze_structure(
        self,
        file_content: bytes,
        file_name: str,
    ) -> dict:
        """深度分析表格結構，回傳完整的欄位位置地圖 (Field Position Map)"""
        file_type = file_name.split('.')[-1].lower()
        field_map = await self._analyzer.analyze(file_content, file_name)
        return {
            "success": True,
            "file_type": file_type,
            "field_map": field_map,
            "total_fields": len(field_map),
        }

    async def _deep_analyze_excel(self, content: bytes) -> list[dict]:
        """向後相容：委派給 autofill_core。"""
        return await self._analyzer.analyze_excel(content)

    async def _deep_analyze_word(self, content: bytes) -> list[dict]:
        """向後相容：委派給 autofill_core。"""
        return await self._analyzer.analyze_word(content)

    async def _extract_excel_text(self, content: bytes) -> str:
        """向後相容：委派給 autofill_core。"""
        return await self._analyzer._extract_excel_text(content)

    async def _extract_word_text(self, content: bytes) -> str:
        """向後相容：委派給 autofill_core。"""
        return await self._analyzer._extract_word_text(content)

    # ================================================================
    # AI 欄位映射
    # ================================================================

    async def ai_map_fields(
        self,
        field_map: list[dict],
        inspection_results: list[dict],
    ) -> dict:
        """使用 AI 將深度分析的欄位地圖與檢查結果進行智慧映射"""
        field_summary = []
        for f in field_map:
            field_summary.append({
                "field_id": f["field_id"],
                "field_name": f["field_name"],
                "field_type": f["field_type"],
            })

        results_summary = []
        for idx, result in enumerate(inspection_results):
            item = {
                "index": idx,
                "equipment_name": result.get("equipment_name", ""),
                "equipment_type": result.get("equipment_type", ""),
                "condition": result.get("condition_assessment", ""),
                "is_anomaly": result.get("is_anomaly", False),
                "readings": result.get("extracted_values", {}),
                "anomaly": result.get("anomaly_description", ""),
                "notes": result.get("notes", ""),
            }
            results_summary.append(item)

        try:
            model = genai.GenerativeModel(settings.gemini_flash_model)

            prompt = f"""你是一位工業定檢表單自動填寫專家。請將 AI 檢查結果映射到定檢表格欄位。

【表格欄位】
{json.dumps(field_summary, ensure_ascii=False, indent=2)}

【AI 檢查結果】
{json.dumps(results_summary, ensure_ascii=False, indent=2)}

請為每個表格欄位建議要填入的值。回傳 JSON 陣列，格式如下：
[
  {{
    "field_id": "欄位ID",
    "suggested_value": "建議填入的值",
    "source": "來源說明（哪個檢查結果的哪個欄位）",
    "confidence": 0.95
  }},
  ...
]

規則：
1. 日期欄位填入檢查日期
2. 數值欄位從 readings 中匹配相關讀數
3. 狀態/判定欄位填入「合格」或「不合格」（根據 is_anomaly）
4. 文字欄位填入對應的描述文字
5. 無法映射的欄位不要包含在結果中
6. confidence 為 0-1 之間的數值，表示映射的信心程度

只回應 JSON 陣列，不要其他文字。"""

            response = model.generate_content(prompt)

            json_match = re.search(r'\[[\s\S]*\]', response.text)
            if json_match:
                mappings = json.loads(json_match.group())
                return {
                    "success": True,
                    "mappings": mappings,
                    "unmapped_fields": [
                        f["field_id"] for f in field_map
                        if f["field_id"] not in {m["field_id"] for m in mappings}
                    ],
                }

        except Exception as e:
            logger.error(f"AI map fields failed: {e}")

        return {
            "success": False,
            "mappings": [],
            "unmapped_fields": [f["field_id"] for f in field_map],
            "error": "AI 映射失敗，請手動設定",
        }

    async def save_field_mappings(
        self,
        template: dict,
        mappings: dict[str, str]
    ):
        """儲存欄位對應設定（直接操作 template dict）"""
        for field in template.get("fields", []):
            if field["field_id"] in mappings:
                field["mapping"] = mappings[field["field_id"]]

    # ================================================================
    # 舊版解析（保留相容性）
    # ================================================================

    async def _parse_excel_template(self, content: bytes) -> tuple[list[dict], str]:
        """解析 Excel 模板（舊版，保留相容性）"""
        fields = []
        structure_lines = []

        wb = load_workbook(io.BytesIO(content))
        ws = wb.active

        for row_idx, row in enumerate(ws.iter_rows(max_row=50), 1):
            for cell in row:
                if cell.value:
                    value = str(cell.value).strip()
                    structure_lines.append(f"{cell.coordinate}: {value}")

                    if any(kw in value for kw in [':', '：', '日期', '姓名', '編號', '設備', '檢查', '備註']):
                        fields.append({
                            "field_id": f"field_{cell.coordinate}",
                            "field_name": value.rstrip(':：'),
                            "field_type": guess_field_type(value),
                            "location": cell.coordinate,
                            "mapping": None,
                        })

        return fields, "\n".join(structure_lines[:100])

    async def _parse_word_template(self, content: bytes) -> tuple[list[dict], str]:
        """解析 Word 模板（舊版，保留相容性）"""
        fields = []
        structure_lines = []

        doc = Document(io.BytesIO(content))

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                structure_lines.append(f"Para {para_idx}: {text}")

                if any(kw in text for kw in [':', '：', '____', '＿＿']):
                    fields.append({
                        "field_id": f"field_para_{para_idx}",
                        "field_name": text.split(':')[0].split('：')[0].strip(),
                        "field_type": "text",
                        "location": f"paragraph_{para_idx}",
                        "mapping": None,
                    })

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        loc = f"table_{table_idx}_r{row_idx}_c{cell_idx}"
                        structure_lines.append(f"{loc}: {text}")

        return fields, "\n".join(structure_lines[:100])

    async def _ai_suggest_mappings(
        self,
        fields: list[dict],
        raw_structure: str
    ) -> dict[str, str]:
        """使用 AI 建議欄位對應"""
        try:
            model = genai.GenerativeModel(settings.gemini_flash_model)

            fields_text = "\n".join([
                f"- {f['field_id']}: {f['field_name']} ({f['field_type']})"
                for f in fields
            ])

            prompt = f"""
你是一位表單分析專家。請分析以下廠商報告模板的欄位，並建議對應到巡檢資料的欄位。

【模板欄位】
{fields_text}

【巡檢資料可用欄位】
- equipment_name: 設備名稱
- equipment_type: 設備類型
- equipment_id: 設備編號
- inspection_date: 巡檢日期
- inspector_name: 巡檢人員
- location: 位置
- condition_assessment: 狀況評估
- anomaly_description: 異常描述
- is_anomaly: 是否異常
- extracted_values: 儀表讀數（dict）
- notes: 備註

請以 JSON 格式回應，格式如下：
{{"field_id": "對應的巡檢欄位", ...}}

只回應 JSON，不要其他文字。無法對應的欄位請填 null。
"""

            response = model.generate_content(prompt)

            json_match = re.search(r'\{[^{}]+\}', response.text, re.DOTALL)
            if json_match:
                mappings = json.loads(json_match.group())
                return {k: v for k, v in mappings.items() if v}

        except Exception as e:
            logger.error(f"AI suggest mappings failed: {e}")

        return {}
