"""
表單填入服務 - AI 模板分析與報告生成

支援深度表格結構解析、AI 智慧欄位映射、自動回填引擎
"""

import logging
import uuid
import io
import json
import re
import copy
from typing import Optional
from datetime import datetime

import google.generativeai as genai
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter, column_index_from_string
from docx import Document

from app.config import settings

logger = logging.getLogger(__name__)


# ============ 定檢資料可用欄位定義 ============

INSPECTION_FIELDS = {
    "equipment_name": {"label": "設備名稱", "type": "text"},
    "equipment_type": {"label": "設備類型", "type": "text"},
    "equipment_id": {"label": "設備編號", "type": "text"},
    "inspection_date": {"label": "檢查日期", "type": "date"},
    "inspector_name": {"label": "檢查人員", "type": "text"},
    "location": {"label": "位置/廠區", "type": "text"},
    "condition_assessment": {"label": "狀況評估", "type": "text"},
    "anomaly_description": {"label": "異常描述", "type": "text"},
    "is_anomaly": {"label": "是否異常", "type": "checkbox"},
    "notes": {"label": "備註", "type": "text"},
    # 量測讀數（動態 key，從 extracted_values 展開）
    "extracted_values": {"label": "儀表讀數/量測值", "type": "dict"},
}

# 欄位標籤關鍵字 — 用於自動偵測
FIELD_KEYWORDS = [
    ':', '：', '日期', '姓名', '編號', '設備', '檢查', '備註',
    '人員', '地點', '位置', '廠區', '型號', '規格', '狀態', '狀況',
    '結果', '判定', '溫度', '壓力', '電流', '電壓', '轉速', '流量',
    '讀數', '數值', '合格', '不合格', '正常', '異常', '測量',
    '頻率', '振動', '噪音', '油位', '水位', '濕度',
]


class FormFillService:
    """表單自動填入服務"""

    def __init__(self):
        genai.configure(api_key=settings.gemini_api_key)

        # TODO: 正式環境改用資料庫
        self._templates: dict[str, dict] = {}
        self._reports: dict[str, dict] = {}

    # ================================================================
    # 模板 CRUD
    # ================================================================

    async def list_templates(self) -> list[dict]:
        """列出所有模板"""
        return list(self._templates.values())

    async def get_template(self, template_id: str) -> Optional[dict]:
        """取得單一模板"""
        return self._templates.get(template_id)

    async def delete_template(self, template_id: str):
        """刪除模板"""
        if template_id in self._templates:
            del self._templates[template_id]

    # ================================================================
    # 模板分析（上傳時）
    # ================================================================

    async def analyze_template(
        self,
        file_content: bytes,
        file_name: str,
        vendor_name: str,
        template_name: str,
        description: Optional[str] = None
    ) -> dict:
        """
        使用 AI 分析模板結構

        自動識別欄位並建議與巡檢資料的對應關係
        """
        template_id = str(uuid.uuid4())
        file_type = file_name.split('.')[-1].lower()

        # 解析模板內容
        if file_type == 'xlsx':
            fields, raw_structure = await self._parse_excel_template(file_content)
        elif file_type == 'docx':
            fields, raw_structure = await self._parse_word_template(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # 使用 AI 建議欄位對應
        suggested_mappings = await self._ai_suggest_mappings(fields, raw_structure)

        # 儲存模板
        template = {
            "id": template_id,
            "name": template_name,
            "vendor_name": vendor_name,
            "file_type": file_type,
            "description": description,
            "fields": fields,
            "file_content": file_content,  # 保存原始檔案
            "created_at": datetime.now().isoformat(),
        }
        self._templates[template_id] = template

        return {
            "success": True,
            "template_id": template_id,
            "detected_fields": fields,
            "suggested_mappings": suggested_mappings,
            "message": f"成功分析模板，識別到 {len(fields)} 個欄位"
        }

    # ================================================================
    # 深度結構分析 (新增)
    # ================================================================

    async def analyze_structure(
        self,
        file_content: bytes,
        file_name: str,
    ) -> dict:
        """
        深度分析表格結構，回傳完整的欄位位置地圖 (Field Position Map)

        不儲存模板，僅回傳分析結果供前端預覽。
        """
        file_type = file_name.split('.')[-1].lower()

        if file_type == 'xlsx':
            field_map = await self._deep_analyze_excel(file_content)
        elif file_type == 'docx':
            field_map = await self._deep_analyze_word(file_content)
        else:
            raise ValueError(f"不支援的檔案格式: {file_type}")

        return {
            "success": True,
            "file_type": file_type,
            "field_map": field_map,
            "total_fields": len(field_map),
        }

    async def _deep_analyze_excel(self, content: bytes) -> list[dict]:
        """
        深度分析 Excel 表格結構

        回傳每個可填入欄位的位置資訊，包含：
        - 欄位標籤文字與位置
        - 對應的值儲存格座標 (value_cell)
        - 合併儲存格處理
        - 欄位類型推測
        """
        fields = []
        wb = load_workbook(io.BytesIO(content))

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # 收集合併儲存格資訊
            merged_ranges = list(ws.merged_cells.ranges)

            # 建立合併儲存格查找表: cell_coord -> merge_range
            merge_lookup = {}
            for mr in merged_ranges:
                for row in range(mr.min_row, mr.max_row + 1):
                    for col in range(mr.min_col, mr.max_col + 1):
                        coord = f"{get_column_letter(col)}{row}"
                        merge_lookup[coord] = {
                            "range": str(mr),
                            "top_left": f"{get_column_letter(mr.min_col)}{mr.min_row}",
                            "rows": mr.max_row - mr.min_row + 1,
                            "cols": mr.max_col - mr.min_col + 1,
                        }

            max_row = min(ws.max_row or 1, 200)
            max_col = min(ws.max_column or 1, 50)

            for row_idx in range(1, max_row + 1):
                for col_idx in range(1, max_col + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if cell.value is None:
                        continue

                    value = str(cell.value).strip()
                    if not value:
                        continue

                    coord = f"{get_column_letter(col_idx)}{row_idx}"

                    # 跳過合併儲存格中非左上角的儲存格
                    if coord in merge_lookup and merge_lookup[coord]["top_left"] != coord:
                        continue

                    # 判斷是否為欄位標籤
                    is_label = self._is_field_label(value)
                    if not is_label:
                        continue

                    # 找到對應的值儲存格（標籤右邊或下方的空白/可填入儲存格）
                    value_cell = self._find_value_cell_excel(
                        ws, row_idx, col_idx, max_row, max_col, merge_lookup
                    )

                    field_type = self._guess_field_type(value)

                    fields.append({
                        "field_id": f"excel_{sheet_name}_{coord}",
                        "field_name": value.rstrip(':：ˍ_ '),
                        "field_type": field_type,
                        "label_location": {
                            "sheet": sheet_name,
                            "cell": coord,
                            "row": row_idx,
                            "column": col_idx,
                        },
                        "value_location": value_cell,
                        "is_merged": coord in merge_lookup,
                        "merge_info": merge_lookup.get(coord),
                        "mapping": None,
                    })

        return fields

    def _find_value_cell_excel(
        self, ws, label_row: int, label_col: int,
        max_row: int, max_col: int,
        merge_lookup: dict,
    ) -> Optional[dict]:
        """
        找到標籤欄位對應的值儲存格

        策略：
        1. 先檢查右邊的儲存格
        2. 再檢查下方的儲存格
        3. 考慮合併儲存格
        """
        sheet_name = ws.title

        # 策略 1: 標籤右邊
        for offset in range(1, 4):
            next_col = label_col + offset
            if next_col > max_col:
                break
            next_coord = f"{get_column_letter(next_col)}{label_row}"
            next_cell = ws.cell(row=label_row, column=next_col)

            # 如果右邊的儲存格是空的或包含佔位符，就是值儲存格
            cell_val = next_cell.value
            if cell_val is None or self._is_placeholder(str(cell_val)):
                return {
                    "sheet": sheet_name,
                    "cell": next_coord,
                    "row": label_row,
                    "column": next_col,
                    "direction": "right",
                    "offset": offset,
                }

            # 如果右邊儲存格有值但不是標籤，可能也是值
            if not self._is_field_label(str(cell_val)):
                return {
                    "sheet": sheet_name,
                    "cell": next_coord,
                    "row": label_row,
                    "column": next_col,
                    "direction": "right",
                    "offset": offset,
                }

        # 策略 2: 標籤下方
        for offset in range(1, 3):
            next_row = label_row + offset
            if next_row > max_row:
                break
            next_coord = f"{get_column_letter(label_col)}{next_row}"
            next_cell = ws.cell(row=next_row, column=label_col)

            cell_val = next_cell.value
            if cell_val is None or self._is_placeholder(str(cell_val)):
                return {
                    "sheet": sheet_name,
                    "cell": next_coord,
                    "row": next_row,
                    "column": label_col,
                    "direction": "below",
                    "offset": offset,
                }

        return None

    async def _deep_analyze_word(self, content: bytes) -> list[dict]:
        """
        深度分析 Word 文件中的表格結構
        """
        fields = []
        doc = Document(io.BytesIO(content))

        # 分析段落中的欄位
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            if self._is_field_label(text) or '____' in text or '＿＿' in text:
                field_name = text.split(':')[0].split('：')[0].strip()
                field_name = field_name.rstrip('_＿ ')

                fields.append({
                    "field_id": f"word_para_{para_idx}",
                    "field_name": field_name,
                    "field_type": self._guess_field_type(field_name),
                    "label_location": {
                        "type": "paragraph",
                        "paragraph_index": para_idx,
                    },
                    "value_location": {
                        "type": "paragraph",
                        "paragraph_index": para_idx,
                        "replace_pattern": "after_colon",  # 替換冒號後的內容
                    },
                    "mapping": None,
                })

        # 分析表格中的欄位
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue

                    if not self._is_field_label(text):
                        continue

                    field_name = text.rstrip(':：_＿ ')

                    # 找對應的值儲存格（右邊或下方）
                    value_loc = self._find_value_cell_word(
                        table, row_idx, cell_idx
                    )

                    fields.append({
                        "field_id": f"word_t{table_idx}_r{row_idx}_c{cell_idx}",
                        "field_name": field_name,
                        "field_type": self._guess_field_type(field_name),
                        "label_location": {
                            "type": "table",
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "cell_index": cell_idx,
                        },
                        "value_location": value_loc,
                        "mapping": None,
                    })

        return fields

    def _find_value_cell_word(
        self, table, label_row: int, label_col: int
    ) -> Optional[dict]:
        """找到 Word 表格中標籤對應的值儲存格"""
        rows = table.rows
        num_cols = len(rows[label_row].cells) if label_row < len(rows) else 0

        # 策略 1: 右邊的儲存格
        if label_col + 1 < num_cols:
            right_cell = rows[label_row].cells[label_col + 1]
            right_text = right_cell.text.strip()
            if not right_text or self._is_placeholder(right_text):
                return {
                    "type": "table",
                    "table_index": None,  # 由呼叫者填入
                    "row_index": label_row,
                    "cell_index": label_col + 1,
                    "direction": "right",
                }

        # 策略 2: 下方的儲存格
        if label_row + 1 < len(rows):
            below_cell = rows[label_row + 1].cells[label_col]
            below_text = below_cell.text.strip()
            if not below_text or self._is_placeholder(below_text):
                return {
                    "type": "table",
                    "table_index": None,
                    "row_index": label_row + 1,
                    "cell_index": label_col,
                    "direction": "below",
                }

        return None

    # ================================================================
    # AI 欄位映射
    # ================================================================

    async def ai_map_fields(
        self,
        field_map: list[dict],
        inspection_results: list[dict],
    ) -> dict:
        """
        使用 AI 將深度分析的欄位地圖與檢查結果進行智慧映射

        Args:
            field_map: 來自 analyze_structure 的欄位位置地圖
            inspection_results: AI 分析後的檢查結果列表

        Returns:
            映射結果，包含每個欄位建議填入的值與信心度
        """
        field_summary = []
        for f in field_map:
            field_summary.append({
                "field_id": f["field_id"],
                "field_name": f["field_name"],
                "field_type": f["field_type"],
            })

        # 組合所有檢查結果為文字
        results_summary = []
        for idx, result in enumerate(inspection_results):
            item = {
                "index": idx,
                "equipment_name": result.get("equipment_name", ""),
                "equipment_type": result.get("equipment_type", ""),
                "condition": result.get("condition_assessment", ""),
                "is_anomaly": result.get("is_anomaly", False),
                "readings": result.get("extracted_values", {}),
                "anomaly": result.get("anomaly_description", ""),
                "notes": result.get("notes", ""),
            }
            results_summary.append(item)

        try:
            model = genai.GenerativeModel('gemini-2.0-flash')

            prompt = f"""你是一位工業定檢表單自動填寫專家。請將 AI 檢查結果映射到定檢表格欄位。

【表格欄位】
{json.dumps(field_summary, ensure_ascii=False, indent=2)}

【AI 檢查結果】
{json.dumps(results_summary, ensure_ascii=False, indent=2)}

請為每個表格欄位建議要填入的值。回傳 JSON 陣列，格式如下：
[
  {{
    "field_id": "欄位ID",
    "suggested_value": "建議填入的值",
    "source": "來源說明（哪個檢查結果的哪個欄位）",
    "confidence": 0.95
  }},
  ...
]

規則：
1. 日期欄位填入檢查日期
2. 數值欄位從 readings 中匹配相關讀數
3. 狀態/判定欄位填入「合格」或「不合格」（根據 is_anomaly）
4. 文字欄位填入對應的描述文字
5. 無法映射的欄位不要包含在結果中
6. confidence 為 0-1 之間的數值，表示映射的信心程度

只回應 JSON 陣列，不要其他文字。"""

            response = model.generate_content(prompt)

            # 解析 JSON 回應
            json_match = re.search(r'\[[\s\S]*\]', response.text)
            if json_match:
                mappings = json.loads(json_match.group())
                return {
                    "success": True,
                    "mappings": mappings,
                    "unmapped_fields": [
                        f["field_id"] for f in field_map
                        if f["field_id"] not in {m["field_id"] for m in mappings}
                    ],
                }

        except Exception as e:
            logger.error(f"AI map fields failed: {e}")

        return {
            "success": False,
            "mappings": [],
            "unmapped_fields": [f["field_id"] for f in field_map],
            "error": "AI 映射失敗，請手動設定",
        }

    async def save_field_mappings(
        self,
        template_id: str,
        mappings: dict[str, str]
    ):
        """儲存欄位對應設定"""
        if template_id not in self._templates:
            raise ValueError(f"Template not found: {template_id}")

        template = self._templates[template_id]
        for field in template["fields"]:
            if field["field_id"] in mappings:
                field["mapping"] = mappings[field["field_id"]]

    # ================================================================
    # 自動回填引擎 (新增)
    # ================================================================

    async def auto_fill(
        self,
        file_content: bytes,
        file_name: str,
        field_map: list[dict],
        fill_values: list[dict],
    ) -> bytes:
        """
        執行自動回填：將值寫入原始文件的指定位置

        Args:
            file_content: 原始文件內容
            file_name: 文件名稱（判斷格式用）
            field_map: 欄位位置地圖（含 value_location）
            fill_values: 要填入的值列表
                [{"field_id": "...", "value": "..."}, ...]

        Returns:
            回填後的文件 bytes
        """
        file_type = file_name.split('.')[-1].lower()

        # 建立 field_id -> value 的查找表
        value_lookup = {fv["field_id"]: fv["value"] for fv in fill_values}

        # 建立 field_id -> field_map entry 的查找表
        field_lookup = {f["field_id"]: f for f in field_map}

        if file_type == 'xlsx':
            return await self._auto_fill_excel(
                file_content, field_lookup, value_lookup
            )
        elif file_type == 'docx':
            return await self._auto_fill_word(
                file_content, field_lookup, value_lookup
            )
        else:
            raise ValueError(f"不支援的檔案格式: {file_type}")

    async def _auto_fill_excel(
        self,
        file_content: bytes,
        field_lookup: dict,
        value_lookup: dict,
    ) -> bytes:
        """
        回填 Excel 檔案

        保留原始格式（字體、邊框、合併儲存格、樣式等）
        """
        wb = load_workbook(io.BytesIO(file_content))

        for field_id, value in value_lookup.items():
            field = field_lookup.get(field_id)
            if not field:
                continue

            val_loc = field.get("value_location")
            if not val_loc:
                # 沒有找到值儲存格，嘗試寫入標籤儲存格
                label_loc = field.get("label_location", {})
                sheet_name = label_loc.get("sheet")
                cell_coord = label_loc.get("cell")
                if sheet_name and cell_coord and sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    ws[cell_coord] = value
                continue

            sheet_name = val_loc.get("sheet")
            cell_coord = val_loc.get("cell")

            if not sheet_name or not cell_coord:
                continue

            if sheet_name not in wb.sheetnames:
                logger.warning(f"Sheet '{sheet_name}' not found, skipping field {field_id}")
                continue

            ws = wb[sheet_name]

            # 根據欄位類型轉換值
            typed_value = self._convert_value(value, field.get("field_type", "text"))

            # 寫入值，保留原始格式
            target_cell = ws[cell_coord]

            # 複製原始格式資訊
            original_font = copy.copy(target_cell.font) if target_cell.font else None
            original_alignment = copy.copy(target_cell.alignment) if target_cell.alignment else None
            original_number_format = target_cell.number_format

            target_cell.value = typed_value

            # 還原格式
            if original_font:
                target_cell.font = original_font
            if original_alignment:
                target_cell.alignment = original_alignment
            if original_number_format:
                target_cell.number_format = original_number_format

        # 輸出為 bytes
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output.read()

    async def _auto_fill_word(
        self,
        file_content: bytes,
        field_lookup: dict,
        value_lookup: dict,
    ) -> bytes:
        """
        回填 Word 檔案

        保留原始格式（字體、段落樣式等）
        """
        doc = Document(io.BytesIO(file_content))

        for field_id, value in value_lookup.items():
            field = field_lookup.get(field_id)
            if not field:
                continue

            val_loc = field.get("value_location")
            if not val_loc:
                continue

            loc_type = val_loc.get("type")

            if loc_type == "paragraph":
                para_idx = val_loc.get("paragraph_index")
                if para_idx is not None and para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    replace_pattern = val_loc.get("replace_pattern", "after_colon")

                    if replace_pattern == "after_colon":
                        # 替換冒號/：後面的內容
                        text = para.text
                        for sep in ['：', ':']:
                            if sep in text:
                                prefix = text.split(sep)[0] + sep
                                self._replace_paragraph_text_preserve_format(
                                    para, f"{prefix} {value}"
                                )
                                break
                    else:
                        # 替換佔位符
                        self._replace_paragraph_text_preserve_format(para, value)

            elif loc_type == "table":
                table_idx = val_loc.get("table_index")
                row_idx = val_loc.get("row_index")
                cell_idx = val_loc.get("cell_index")

                if (table_idx is not None and
                        table_idx < len(doc.tables) and
                        row_idx is not None and
                        cell_idx is not None):
                    table = doc.tables[table_idx]
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        if cell_idx < len(row.cells):
                            cell = row.cells[cell_idx]
                            # 保留格式寫入
                            if cell.paragraphs:
                                self._replace_paragraph_text_preserve_format(
                                    cell.paragraphs[0], str(value)
                                )
                            else:
                                cell.text = str(value)

        # 輸出為 bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def _replace_paragraph_text_preserve_format(self, paragraph, new_text: str):
        """替換段落文字但保留第一個 run 的格式"""
        if not paragraph.runs:
            paragraph.text = new_text
            return

        # 保留第一個 run 的格式
        first_run = paragraph.runs[0]

        # 清除所有 runs
        for run in paragraph.runs:
            run.text = ""

        # 在第一個 run 中設定新文字
        first_run.text = new_text

    # ================================================================
    # 預覽回填
    # ================================================================

    async def preview_fill(
        self,
        template_id: str,
        inspection_data: dict
    ) -> dict:
        """預覽填入結果（舊版 API 相容）"""
        template = self._templates.get(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        field_values = {}
        warnings = []

        for field in template["fields"]:
            mapping = field.get("mapping")
            if mapping and mapping in inspection_data:
                value = inspection_data[mapping]
                if isinstance(value, dict):
                    value = str(value)
                field_values[field["field_name"]] = str(value) if value else ""
            else:
                field_values[field["field_name"]] = ""
                if mapping:
                    warnings.append(f"欄位 '{field['field_name']}' 對應的資料不存在")

        return {
            "template_name": template["name"],
            "vendor_name": template["vendor_name"],
            "field_values": field_values,
            "warnings": warnings,
        }

    async def preview_auto_fill(
        self,
        field_map: list[dict],
        fill_values: list[dict],
    ) -> dict:
        """
        預覽自動回填結果（新版 API）

        回傳每個欄位即將填入的值以及信心度標記
        """
        value_lookup = {fv["field_id"]: fv for fv in fill_values}

        preview_items = []
        warnings = []

        for field in field_map:
            fid = field["field_id"]
            fv = value_lookup.get(fid)

            item = {
                "field_id": fid,
                "field_name": field["field_name"],
                "field_type": field.get("field_type", "text"),
                "value": fv["value"] if fv else None,
                "confidence": fv.get("confidence", 0.0) if fv else 0.0,
                "source": fv.get("source", "") if fv else "",
                "has_target": field.get("value_location") is not None,
            }
            preview_items.append(item)

            if not fv:
                warnings.append(f"欄位 '{field['field_name']}' 無對應值")
            elif item["confidence"] < 0.7:
                warnings.append(
                    f"欄位 '{field['field_name']}' 映射信心度較低 ({item['confidence']:.0%})，建議確認"
                )
            if not item["has_target"]:
                warnings.append(
                    f"欄位 '{field['field_name']}' 找不到值儲存格位置，無法回填"
                )

        return {
            "preview_items": preview_items,
            "total_fields": len(field_map),
            "filled_count": sum(1 for p in preview_items if p["value"] is not None),
            "warnings": warnings,
        }

    # ================================================================
    # 報告生成（舊版 API 相容）
    # ================================================================

    async def generate_report(
        self,
        report_id: str,
        template_id: str,
        inspection_data: dict,
        output_format: str = "xlsx"
    ):
        """產生報告"""
        template = self._templates.get(template_id)
        if not template:
            raise ValueError(f"Template not found: {template_id}")

        try:
            # 根據模板類型處理
            if template["file_type"] == "xlsx":
                output_path = await self._fill_excel(template, inspection_data, report_id)
            elif template["file_type"] == "docx":
                output_path = await self._fill_word(template, inspection_data, report_id)
            else:
                raise ValueError(f"Unsupported template type: {template['file_type']}")

            # 更新報告狀態
            self._reports[report_id] = {
                "id": report_id,
                "status": "completed",
                "template_id": template_id,
                "output_path": output_path,
                "created_at": datetime.now().isoformat(),
            }

        except Exception as e:
            logger.error(f"Generate report failed: {e}")
            self._reports[report_id] = {
                "id": report_id,
                "status": "failed",
                "error": str(e),
            }
            raise

    async def _fill_excel(
        self,
        template: dict,
        inspection_data: dict,
        report_id: str
    ) -> str:
        """填入 Excel 模板"""
        wb = load_workbook(io.BytesIO(template["file_content"]))
        ws = wb.active

        for field in template["fields"]:
            mapping = field.get("mapping")
            if mapping and mapping in inspection_data:
                value = inspection_data[mapping]

                # 找到要填入的儲存格 (欄位位置的右邊或下方)
                location = field["location"]
                # 簡化：假設填入同一儲存格
                # 實際應用需要更複雜的邏輯
                ws[location] = value

        # 儲存
        output_path = f"/tmp/report_{report_id}.xlsx"
        wb.save(output_path)

        return output_path

    async def _fill_word(
        self,
        template: dict,
        inspection_data: dict,
        report_id: str
    ) -> str:
        """填入 Word 模板"""
        doc = Document(io.BytesIO(template["file_content"]))

        # 替換文字中的佔位符
        for para in doc.paragraphs:
            for field in template["fields"]:
                mapping = field.get("mapping")
                if mapping and mapping in inspection_data:
                    value = str(inspection_data[mapping] or "")
                    # 替換格式: {{field_name}}
                    placeholder = f"{{{{{field['field_name']}}}}}"
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, value)

        output_path = f"/tmp/report_{report_id}.docx"
        doc.save(output_path)

        return output_path

    async def get_report_status(self, report_id: str) -> Optional[dict]:
        """取得報告狀態"""
        report = self._reports.get(report_id)
        if not report:
            return None

        return {
            "success": report["status"] == "completed",
            "report_id": report_id,
            "status": report["status"],
            "message": "報告已完成" if report["status"] == "completed" else report.get("error", "處理中"),
            "download_url": f"/api/reports/{report_id}/download" if report["status"] == "completed" else None,
        }

    async def get_report_file(self, report_id: str) -> Optional[str]:
        """取得報告檔案路徑"""
        report = self._reports.get(report_id)
        if report and report.get("status") == "completed":
            return report.get("output_path")
        return None

    # ================================================================
    # 輔助方法
    # ================================================================

    async def _parse_excel_template(self, content: bytes) -> tuple[list[dict], str]:
        """解析 Excel 模板（舊版，保留相容性）"""
        fields = []
        structure_lines = []

        wb = load_workbook(io.BytesIO(content))
        ws = wb.active

        for row_idx, row in enumerate(ws.iter_rows(max_row=50), 1):
            for cell in row:
                if cell.value:
                    # 檢測可能的欄位標籤
                    value = str(cell.value).strip()
                    structure_lines.append(f"{cell.coordinate}: {value}")

                    # 簡單規則：包含冒號或特定關鍵字的可能是欄位標籤
                    if any(kw in value for kw in [':', '：', '日期', '姓名', '編號', '設備', '檢查', '備註']):
                        fields.append({
                            "field_id": f"field_{cell.coordinate}",
                            "field_name": value.rstrip(':：'),
                            "field_type": self._guess_field_type(value),
                            "location": cell.coordinate,
                            "mapping": None,
                        })

        return fields, "\n".join(structure_lines[:100])

    async def _parse_word_template(self, content: bytes) -> tuple[list[dict], str]:
        """解析 Word 模板（舊版，保留相容性）"""
        fields = []
        structure_lines = []

        doc = Document(io.BytesIO(content))

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                structure_lines.append(f"Para {para_idx}: {text}")

                # 檢測可能的欄位
                if any(kw in text for kw in [':', '：', '____', '＿＿']):
                    fields.append({
                        "field_id": f"field_para_{para_idx}",
                        "field_name": text.split(':')[0].split('：')[0].strip(),
                        "field_type": "text",
                        "location": f"paragraph_{para_idx}",
                        "mapping": None,
                    })

        # 也檢查表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        loc = f"table_{table_idx}_r{row_idx}_c{cell_idx}"
                        structure_lines.append(f"{loc}: {text}")

        return fields, "\n".join(structure_lines[:100])

    def _is_field_label(self, text: str) -> bool:
        """判斷文字是否為欄位標籤"""
        text = text.strip()
        if not text or len(text) > 50:
            return False
        return any(kw in text for kw in FIELD_KEYWORDS)

    def _is_placeholder(self, text: str) -> bool:
        """判斷文字是否為佔位符"""
        text = text.strip()
        if not text:
            return True
        # 常見佔位符模式
        placeholder_patterns = [
            r'^[_＿]{2,}$',     # 底線
            r'^\{\{.*\}\}$',    # {{placeholder}}
            r'^<.*>$',          # <placeholder>
            r'^\[.*\]$',        # [placeholder]
            r'^/{2,}$',         # ///
            r'^\s+$',           # 純空白
        ]
        return any(re.match(p, text) for p in placeholder_patterns)

    def _guess_field_type(self, field_name: str) -> str:
        """猜測欄位類型"""
        name_lower = field_name.lower()

        if any(kw in name_lower for kw in ['日期', 'date', '時間', 'time']):
            return 'date'
        elif any(kw in name_lower for kw in [
            '數量', '數值', 'number', '金額', '溫度', '壓力',
            '電流', '電壓', '轉速', '流量', '讀數', '頻率',
            '振動', '噪音', '油位', '水位', '濕度',
        ]):
            return 'number'
        elif any(kw in name_lower for kw in [
            '是否', '確認', 'check', '合格', '判定', '正常', '異常',
        ]):
            return 'checkbox'
        else:
            return 'text'

    def _convert_value(self, value, field_type: str):
        """根據欄位類型轉換值"""
        if value is None:
            return None

        if field_type == 'number':
            try:
                if '.' in str(value):
                    return float(value)
                return int(value)
            except (ValueError, TypeError):
                return str(value)
        elif field_type == 'checkbox':
            v = str(value).strip().lower()
            if v in ['true', '1', '是', '合格', '正常', 'yes', 'ok', '通過']:
                return '合格'
            elif v in ['false', '0', '否', '不合格', '異常', 'no', 'ng', '不通過']:
                return '不合格'
            return str(value)
        elif field_type == 'date':
            return str(value)
        else:
            return str(value)

    async def _ai_suggest_mappings(
        self,
        fields: list[dict],
        raw_structure: str
    ) -> dict[str, str]:
        """使用 AI 建議欄位對應"""
        try:
            model = genai.GenerativeModel('gemini-2.0-flash')

            fields_text = "\n".join([
                f"- {f['field_id']}: {f['field_name']} ({f['field_type']})"
                for f in fields
            ])

            prompt = f"""
你是一位表單分析專家。請分析以下廠商報告模板的欄位，並建議對應到巡檢資料的欄位。

【模板欄位】
{fields_text}

【巡檢資料可用欄位】
- equipment_name: 設備名稱
- equipment_type: 設備類型
- equipment_id: 設備編號
- inspection_date: 巡檢日期
- inspector_name: 巡檢人員
- location: 位置
- condition_assessment: 狀況評估
- anomaly_description: 異常描述
- is_anomaly: 是否異常
- extracted_values: 儀表讀數（dict）
- notes: 備註

請以 JSON 格式回應，格式如下：
{{"field_id": "對應的巡檢欄位", ...}}

只回應 JSON，不要其他文字。無法對應的欄位請填 null。
"""

            response = model.generate_content(prompt)

            # 提取 JSON 部分
            json_match = re.search(r'\{[^{}]+\}', response.text, re.DOTALL)
            if json_match:
                mappings = json.loads(json_match.group())
                return {k: v for k, v in mappings.items() if v}

        except Exception as e:
            logger.error(f"AI suggest mappings failed: {e}")

        return {}
