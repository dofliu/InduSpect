"""
表格自動回填功能 — 綜合測試套件

測試目標：
    用不同格式的 Excel/Word 表格，驗證自動回填系統能否正確：
    1. 分析表格結構（欄位偵測）
    2. 推測欄位類型
    3. 找到值儲存格位置
    4. 執行回填並保留格式
    5. 預覽回填結果

測試表格格式：
    Excel:
        A. 簡單 key-value 橫向排列 (標籤在左，值在右)
        B. 垂直排列 (標籤在上，值在下)
        C. 合併儲存格
        D. 多工作表
        E. 混合複雜排列
        F. 含有佔位符的表格
        G. 密集多欄標籤表格
    Word:
        H. 段落式 (冒號分隔)
        I. 表格式 key-value
        J. 混合段落+表格
        K. 底線佔位符段落
"""

import io
import sys
import os
import json
import asyncio
import types
import pytest
from datetime import datetime
from unittest.mock import patch, MagicMock

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches

# 讓 import 找到 app 模組
sys.path.insert(0, os.path.join(os.path.dirname(__file__)))

# Mock 所有缺失的深層依賴，避免 import 失敗
for mod_name in [
    "google.generativeai",
    "sqlalchemy", "sqlalchemy.ext", "sqlalchemy.ext.asyncio",
    "sqlalchemy.orm", "sqlalchemy.dialects", "sqlalchemy.dialects.postgresql",
    "pgvector", "pgvector.sqlalchemy",
    "openai", "httpx", "aiohttp",
]:
    if mod_name not in sys.modules:
        sys.modules[mod_name] = MagicMock()

# Mock settings
os.environ.setdefault("GEMINI_API_KEY", "fake-key-for-test")

from app.services.form_fill import FormFillService, FIELD_KEYWORDS


# ============================================================
# Helper: 建立各種測試用 Excel/Word 檔案
# ============================================================

def create_excel_simple_horizontal() -> bytes:
    """格式 A: 簡單水平 key-value 排列

    | 設備名稱： | (空) | 檢查日期： | (空)    |
    | 設備編號： | (空) | 檢查人員： | (空)    |
    | 溫度讀數： | (空) | 壓力讀數： | (空)    |
    | 判定結果： | (空) | 備註：     | (空)    |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "巡檢表"

    data = [
        ("設備名稱：", None, "檢查日期：", None),
        ("設備編號：", None, "檢查人員：", None),
        ("溫度讀數：", None, "壓力讀數：", None),
        ("判定結果：", None, "備註：", None),
    ]
    for r, row in enumerate(data, 1):
        for c, val in enumerate(row, 1):
            ws.cell(row=r, column=c, value=val)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def create_excel_vertical() -> bytes:
    """格式 B: 垂直排列 (標籤在上，值在下)

    | 設備名稱 | 設備編號 | 檢查日期 | 狀態判定 |
    | (空)     | (空)     | (空)     | (空)     |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "設備清單"

    headers = ["設備名稱", "設備編號", "檢查日期", "狀態判定"]
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
        ws.cell(row=2, column=c, value=None)  # 空白值儲存格

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def create_excel_merged_cells() -> bytes:
    """格式 C: 含合併儲存格

    |   設備基本資訊 (A1:D1 合併)   |
    | 設備名稱： | (空，B2:D2 合併) |
    | 設備編號： | (空)  | 位置：  | (空)  |
    |   檢查結果 (A4:D4 合併)       |
    | 溫度：     | (空)  | 壓力：  | (空)  |
    | 判定結果： | (空，B6:D6 合併) |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "合併表格"

    # 標題合併
    ws.merge_cells('A1:D1')
    ws['A1'] = "設備基本資訊"

    ws['A2'] = "設備名稱："
    ws.merge_cells('B2:D2')  # 值跨多欄

    ws['A3'] = "設備編號："
    ws['C3'] = "位置："

    ws.merge_cells('A4:D4')
    ws['A4'] = "檢查結果"

    ws['A5'] = "溫度："
    ws['C5'] = "壓力："

    ws['A6'] = "判定結果："
    ws.merge_cells('B6:D6')

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def create_excel_multi_sheet() -> bytes:
    """格式 D: 多工作表

    Sheet1 "設備資訊":
        | 設備名稱： | (空) |
        | 設備編號： | (空) |

    Sheet2 "檢查紀錄":
        | 檢查日期： | (空) |
        | 溫度讀數： | (空) |
        | 判定結果： | (空) |
    """
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "設備資訊"
    ws1['A1'] = "設備名稱："
    ws1['A2'] = "設備編號："

    ws2 = wb.create_sheet("檢查紀錄")
    ws2['A1'] = "檢查日期："
    ws2['A2'] = "溫度讀數："
    ws2['A3'] = "判定結果："

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def create_excel_mixed_complex() -> bytes:
    """格式 E: 混合複雜排列（含已填入值、格式等）

    |     定期檢查報告 (A1:F1 合併, 粗體)     |
    | 設備名稱：  | 馬達A | 設備編號： | M-001 | 位置：  | 廠區A  |
    | 檢查日期：  | (空)  | 人員：     | (空)  |
    |     測量數據 (A4:F4 合併)                 |
    | 溫度(°C)：  | (空)  | 壓力(kPa)：| (空)  | 電流(A)：| (空)  |
    | 振動(mm/s)：| (空)  | 噪音(dB)：| (空)  | 轉速(rpm)：| (空)|
    | 判定：      | (空)  | 備註：     | (空，D7:F7 合併)       |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "定期檢查"

    # 標題
    ws.merge_cells('A1:F1')
    ws['A1'] = "定期檢查報告"
    ws['A1'].font = Font(size=16, bold=True)

    # 已有部分填入值
    ws['A2'] = "設備名稱："
    ws['B2'] = "馬達A"
    ws['C2'] = "設備編號："
    ws['D2'] = "M-001"
    ws['E2'] = "位置："
    ws['F2'] = "廠區A"

    ws['A3'] = "檢查日期："
    ws['C3'] = "人員："

    ws.merge_cells('A4:F4')
    ws['A4'] = "測量數據"
    ws['A4'].font = Font(bold=True)

    ws['A5'] = "溫度(°C)："
    ws['C5'] = "壓力(kPa)："
    ws['E5'] = "電流(A)："

    ws['A6'] = "振動(mm/s)："
    ws['C6'] = "噪音(dB)："
    ws['E6'] = "轉速(rpm)："

    ws['A7'] = "判定："
    ws['C7'] = "備註："
    ws.merge_cells('D7:F7')

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def create_excel_with_placeholders() -> bytes:
    """格式 F: 含佔位符

    | 設備名稱： | ________ | 設備編號： | ________ |
    | 檢查日期： | {{date}} | 人員：     | {{name}} |
    | 溫度：     | <填入>   | 壓力：     | [填入]   |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "佔位符表格"

    ws['A1'] = "設備名稱："
    ws['B1'] = "________"
    ws['C1'] = "設備編號："
    ws['D1'] = "________"

    ws['A2'] = "檢查日期："
    ws['B2'] = "{{date}}"
    ws['C2'] = "人員："
    ws['D2'] = "{{name}}"

    ws['A3'] = "溫度："
    ws['B3'] = "<填入>"
    ws['C3'] = "壓力："
    ws['D3'] = "[填入]"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def create_excel_dense_labels() -> bytes:
    """格式 G: 密集多欄標籤（每 2 欄一組 label-value，多行）

    | 設備名稱：| (空) | 型號：    | (空) | 規格：  | (空) |
    | 檢查日期：| (空) | 廠區：    | (空) | 狀態：  | (空) |
    | 溫度：    | (空) | 壓力：    | (空) | 電流：  | (空) |
    | 電壓：    | (空) | 頻率：    | (空) | 濕度：  | (空) |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "密集表格"

    rows = [
        ["設備名稱：", None, "型號：", None, "規格：", None],
        ["檢查日期：", None, "廠區：", None, "狀態：", None],
        ["溫度：", None, "壓力：", None, "電流：", None],
        ["電壓：", None, "頻率：", None, "濕度：", None],
    ]
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row, 1):
            ws.cell(row=r, column=c, value=val)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def create_word_paragraph() -> bytes:
    """格式 H: Word 段落式（冒號分隔）"""
    doc = Document()
    doc.add_heading("設備巡檢報告", level=1)
    doc.add_paragraph("設備名稱：________________")
    doc.add_paragraph("設備編號：________________")
    doc.add_paragraph("檢查日期：________________")
    doc.add_paragraph("檢查人員：________________")
    doc.add_paragraph("溫度讀數：________________")
    doc.add_paragraph("判定結果：________________")
    doc.add_paragraph("備註：________________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_word_table() -> bytes:
    """格式 I: Word 表格式 key-value"""
    doc = Document()
    doc.add_heading("設備檢查紀錄", level=1)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    cells_data = [
        ("設備名稱：", "", "設備編號：", ""),
        ("檢查日期：", "", "檢查人員：", ""),
        ("位置：", "", "廠區：", ""),
        ("溫度：", "", "壓力：", ""),
        ("判定結果：", "", "備註：", ""),
    ]

    for r, row_data in enumerate(cells_data):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_word_mixed() -> bytes:
    """格式 J: Word 混合段落+表格"""
    doc = Document()
    doc.add_heading("綜合巡檢報告", level=1)

    # 段落部分
    doc.add_paragraph("設備名稱：________________")
    doc.add_paragraph("設備編號：________________")

    doc.add_heading("檢查數據", level=2)

    # 表格部分
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    cells_data = [
        ("溫度：", "", "壓力：", ""),
        ("電流：", "", "電壓：", ""),
        ("判定結果：", "", "備註：", ""),
    ]
    for r, row_data in enumerate(cells_data):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_word_underline() -> bytes:
    """格式 K: Word 底線佔位符段落"""
    doc = Document()
    doc.add_heading("檢查記錄表", level=1)
    doc.add_paragraph("設備名稱：＿＿＿＿＿＿＿＿")
    doc.add_paragraph("設備編號：____________________")
    doc.add_paragraph("檢查日期：＿＿＿＿＿＿")
    doc.add_paragraph("檢查人員：____________________")
    doc.add_paragraph("溫度讀數：＿＿＿＿ °C")
    doc.add_paragraph("備註：____________________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# 建立 Service 實例（mock AI）
# ============================================================

def get_service():
    """建立 FormFillService（不需真實 API key）"""
    with patch("google.generativeai.configure"):
        svc = FormFillService()
    return svc


# ============================================================
# 測試用巡檢結果
# ============================================================

SAMPLE_INSPECTION = {
    "equipment_name": "馬達 A-001",
    "equipment_type": "交流馬達",
    "equipment_id": "EQ-2026-001",
    "inspection_date": "2026-03-05",
    "inspector_name": "王小明",
    "location": "B棟 2F",
    "condition_assessment": "運轉正常，輕微振動",
    "anomaly_description": "",
    "is_anomaly": False,
    "extracted_values": {
        "溫度": "72.5",
        "壓力": "101.3",
        "電流": "15.2",
        "電壓": "380",
        "轉速": "1450",
        "頻率": "50",
        "振動": "2.1",
        "噪音": "65",
        "濕度": "45",
    },
    "notes": "下次建議更換軸承",
}


# ============================================================
# 測試結果收集器
# ============================================================

class TestReport:
    """收集測試結果用於報告輸出"""

    def __init__(self):
        self.results = []

    def add(self, test_name: str, format_name: str, passed: bool,
            details: str = "", fields_detected: int = 0,
            fields_expected: int = 0, fill_success: int = 0,
            fill_total: int = 0, issues: list = None):
        self.results.append({
            "test_name": test_name,
            "format_name": format_name,
            "passed": passed,
            "details": details,
            "fields_detected": fields_detected,
            "fields_expected": fields_expected,
            "fill_success": fill_success,
            "fill_total": fill_total,
            "issues": issues or [],
        })


report = TestReport()


# ============================================================
# 測試 1: 輔助方法單元測試
# ============================================================

class TestHelperMethods:
    """測試 _is_field_label, _is_placeholder, _guess_field_type, _convert_value"""

    def setup_method(self):
        self.svc = get_service()

    # ---------- _is_field_label ----------

    def test_label_with_colon(self):
        assert self.svc._is_field_label("設備名稱：") is True
        assert self.svc._is_field_label("Name:") is True

    def test_label_with_keyword(self):
        for kw in ["日期", "檢查", "溫度", "壓力", "電流", "判定", "備註", "設備"]:
            assert self.svc._is_field_label(kw) is True, f"應該識別 '{kw}' 為標籤"

    def test_label_empty_or_long(self):
        assert self.svc._is_field_label("") is False
        assert self.svc._is_field_label("A" * 51) is False

    def test_label_non_keyword(self):
        assert self.svc._is_field_label("Hello World") is False
        assert self.svc._is_field_label("12345") is False

    # ---------- _is_placeholder ----------

    def test_placeholder_underlines(self):
        assert self.svc._is_placeholder("________") is True
        assert self.svc._is_placeholder("＿＿＿＿") is True

    def test_placeholder_template(self):
        assert self.svc._is_placeholder("{{date}}") is True
        assert self.svc._is_placeholder("<填入>") is True
        assert self.svc._is_placeholder("[填入]") is True

    def test_placeholder_empty(self):
        assert self.svc._is_placeholder("") is True
        assert self.svc._is_placeholder("   ") is True

    def test_not_placeholder_real_value(self):
        assert self.svc._is_placeholder("馬達A") is False
        assert self.svc._is_placeholder("72.5") is False

    # ---------- _guess_field_type ----------

    def test_guess_date(self):
        assert self.svc._guess_field_type("檢查日期") == "date"
        assert self.svc._guess_field_type("日期") == "date"

    def test_guess_number(self):
        for name in ["溫度", "壓力", "電流", "電壓", "轉速", "頻率", "振動", "噪音", "濕度"]:
            assert self.svc._guess_field_type(name) == "number", f"'{name}' 應為 number 類型"

    def test_guess_checkbox(self):
        assert self.svc._guess_field_type("判定結果") == "checkbox"
        assert self.svc._guess_field_type("是否異常") == "checkbox"
        assert self.svc._guess_field_type("合格判定") == "checkbox"

    def test_guess_text(self):
        assert self.svc._guess_field_type("設備名稱") == "text"
        assert self.svc._guess_field_type("備註") == "text"

    # ---------- _convert_value ----------

    def test_convert_number_int(self):
        assert self.svc._convert_value("42", "number") == 42

    def test_convert_number_float(self):
        assert self.svc._convert_value("72.5", "number") == 72.5

    def test_convert_number_invalid(self):
        assert self.svc._convert_value("abc", "number") == "abc"

    def test_convert_checkbox_pass(self):
        for v in ["true", "1", "是", "合格", "正常", "yes", "ok", "通過"]:
            assert self.svc._convert_value(v, "checkbox") == "合格", f"'{v}' 應轉為 '合格'"

    def test_convert_checkbox_fail(self):
        for v in ["false", "0", "否", "不合格", "異常", "no", "ng", "不通過"]:
            assert self.svc._convert_value(v, "checkbox") == "不合格", f"'{v}' 應轉為 '不合格'"

    def test_convert_none(self):
        assert self.svc._convert_value(None, "text") is None


# ============================================================
# 測試 2: Excel 結構分析
# ============================================================

class TestExcelStructureAnalysis:
    """測試 _deep_analyze_excel 對不同 Excel 格式的分析能力"""

    def setup_method(self):
        self.svc = get_service()

    @pytest.mark.asyncio
    async def test_format_a_simple_horizontal(self):
        """格式 A: 簡單水平 key-value"""
        content = create_excel_simple_horizontal()
        fields = await self.svc._deep_analyze_excel(content)

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "檢查日期", "設備編號", "檢查人員",
                     "溫度讀數", "壓力讀數", "判定結果", "備註"]

        detected = sum(1 for e in expected if any(e in n for n in names))

        # 記錄測試結果
        issues = [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]
        report.add(
            "結構分析", "A: 簡單水平 Excel",
            detected >= 6,
            f"偵測到 {len(fields)} 個欄位，命中 {detected}/{len(expected)}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 6, f"至少偵測 6 個欄位，實際命中 {detected}: {names}"

        # 檢查值儲存格位置是否正確（應在右邊）
        for f in fields:
            if f.get("value_location"):
                assert f["value_location"]["direction"] in ("right", "below")

    @pytest.mark.asyncio
    async def test_format_b_vertical(self):
        """格式 B: 垂直排列"""
        content = create_excel_vertical()
        fields = await self.svc._deep_analyze_excel(content)

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "設備編號", "檢查日期", "狀態判定"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        # 檢查方向
        below_count = sum(
            1 for f in fields
            if f.get("value_location") and f["value_location"].get("direction") == "below"
        )

        issues = []
        if below_count == 0:
            issues.append("未偵測到 'below' 方向的值儲存格")
        issues += [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "B: 垂直排列 Excel",
            detected >= 3,
            f"偵測到 {len(fields)} 個欄位，命中 {detected}/{len(expected)}，below方向: {below_count}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 3

    @pytest.mark.asyncio
    async def test_format_c_merged_cells(self):
        """格式 C: 合併儲存格"""
        content = create_excel_merged_cells()
        fields = await self.svc._deep_analyze_excel(content)

        names = [f["field_name"] for f in fields]

        # 合併區塊標題不應被視為可填入欄位的值（但會被偵測為標籤）
        actual_fields = [f for f in fields if f.get("value_location")]

        expected_labels = ["設備名稱", "設備編號", "位置", "溫度", "壓力", "判定結果"]
        detected = sum(1 for e in expected_labels if any(e in n for n in names))

        merged_fields = [f for f in fields if f.get("is_merged")]

        issues = [f"未偵測到: {e}" for e in expected_labels if not any(e in n for n in names)]

        report.add(
            "結構分析", "C: 合併儲存格 Excel",
            detected >= 4,
            f"偵測到 {len(fields)} 個欄位（含合併 {len(merged_fields)} 個），有值位置: {len(actual_fields)}",
            fields_detected=len(fields),
            fields_expected=len(expected_labels),
            issues=issues,
        )

        assert detected >= 4

    @pytest.mark.asyncio
    async def test_format_d_multi_sheet(self):
        """格式 D: 多工作表"""
        content = create_excel_multi_sheet()
        fields = await self.svc._deep_analyze_excel(content)

        sheets = set()
        for f in fields:
            loc = f.get("label_location", {})
            if loc.get("sheet"):
                sheets.add(loc["sheet"])

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "設備編號", "檢查日期", "溫度讀數", "判定結果"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        issues = []
        if len(sheets) < 2:
            issues.append(f"僅偵測到 {len(sheets)} 個工作表，期望 2 個")
        issues += [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "D: 多工作表 Excel",
            detected >= 4 and len(sheets) >= 2,
            f"偵測到 {len(fields)} 個欄位，跨 {len(sheets)} 個工作表",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert len(sheets) >= 2, f"應偵測到 2 個工作表, 實際: {sheets}"
        assert detected >= 4

    @pytest.mark.asyncio
    async def test_format_e_mixed_complex(self):
        """格式 E: 混合複雜排列（含已填入值）"""
        content = create_excel_mixed_complex()
        fields = await self.svc._deep_analyze_excel(content)

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "設備編號", "位置", "檢查日期", "人員",
                     "溫度", "壓力", "電流", "振動", "噪音", "轉速", "判定", "備註"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        issues = [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "E: 混合複雜 Excel",
            detected >= 8,
            f"偵測到 {len(fields)} 個欄位，命中 {detected}/{len(expected)}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 8

    @pytest.mark.asyncio
    async def test_format_f_placeholders(self):
        """格式 F: 佔位符表格"""
        content = create_excel_with_placeholders()
        fields = await self.svc._deep_analyze_excel(content)

        # 佔位符儲存格應被識別為值儲存格
        fields_with_target = [f for f in fields if f.get("value_location")]

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "設備編號", "檢查日期", "人員", "溫度", "壓力"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        issues = [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]
        if len(fields_with_target) < detected:
            issues.append(f"僅 {len(fields_with_target)} 個欄位有值位置（期望 {detected}）")

        report.add(
            "結構分析", "F: 佔位符 Excel",
            detected >= 4,
            f"偵測到 {len(fields)} 個欄位，有值位置: {len(fields_with_target)}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 4

    @pytest.mark.asyncio
    async def test_format_g_dense_labels(self):
        """格式 G: 密集多欄標籤"""
        content = create_excel_dense_labels()
        fields = await self.svc._deep_analyze_excel(content)

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "型號", "規格", "檢查日期", "廠區", "狀態",
                     "溫度", "壓力", "電流", "電壓", "頻率", "濕度"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        issues = [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "G: 密集多欄 Excel",
            detected >= 9,
            f"偵測到 {len(fields)} 個欄位，命中 {detected}/{len(expected)}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 9


# ============================================================
# 測試 3: Word 結構分析
# ============================================================

class TestWordStructureAnalysis:
    """測試 _deep_analyze_word 對不同 Word 格式的分析能力"""

    def setup_method(self):
        self.svc = get_service()

    @pytest.mark.asyncio
    async def test_format_h_paragraph(self):
        """格式 H: 段落式（冒號分隔）"""
        content = create_word_paragraph()
        fields = await self.svc._deep_analyze_word(content)

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "設備編號", "檢查日期", "檢查人員", "溫度讀數", "判定結果", "備註"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        para_fields = [f for f in fields if f.get("label_location", {}).get("type") == "paragraph"]

        issues = [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "H: 段落式 Word",
            detected >= 5,
            f"偵測到 {len(fields)} 個欄位（段落: {len(para_fields)}），命中 {detected}/{len(expected)}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 5

    @pytest.mark.asyncio
    async def test_format_i_table(self):
        """格式 I: 表格式 key-value"""
        content = create_word_table()
        fields = await self.svc._deep_analyze_word(content)

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "設備編號", "檢查日期", "檢查人員",
                     "位置", "廠區", "溫度", "壓力", "判定結果", "備註"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        table_fields = [f for f in fields if f.get("label_location", {}).get("type") == "table"]

        issues = [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "I: 表格式 Word",
            detected >= 7,
            f"偵測到 {len(fields)} 個欄位（表格: {len(table_fields)}），命中 {detected}/{len(expected)}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 7

    @pytest.mark.asyncio
    async def test_format_j_mixed(self):
        """格式 J: 混合段落+表格"""
        content = create_word_mixed()
        fields = await self.svc._deep_analyze_word(content)

        names = [f["field_name"] for f in fields]
        para_fields = [f for f in fields if f.get("label_location", {}).get("type") == "paragraph"]
        table_fields = [f for f in fields if f.get("label_location", {}).get("type") == "table"]

        expected = ["設備名稱", "設備編號", "溫度", "壓力", "電流", "電壓", "判定結果", "備註"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        issues = []
        if len(para_fields) == 0:
            issues.append("未偵測到段落欄位")
        if len(table_fields) == 0:
            issues.append("未偵測到表格欄位")
        issues += [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "J: 混合 Word",
            detected >= 5 and len(para_fields) > 0 and len(table_fields) > 0,
            f"偵測到 {len(fields)} 個欄位（段落: {len(para_fields)}, 表格: {len(table_fields)}）",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 5
        assert len(para_fields) > 0, "應偵測到段落欄位"
        assert len(table_fields) > 0, "應偵測到表格欄位"

    @pytest.mark.asyncio
    async def test_format_k_underline(self):
        """格式 K: 底線佔位符段落"""
        content = create_word_underline()
        fields = await self.svc._deep_analyze_word(content)

        names = [f["field_name"] for f in fields]
        expected = ["設備名稱", "設備編號", "檢查日期", "檢查人員", "溫度讀數", "備註"]
        detected = sum(1 for e in expected if any(e in n for n in names))

        issues = [f"未偵測到: {e}" for e in expected if not any(e in n for n in names)]

        report.add(
            "結構分析", "K: 底線佔位符 Word",
            detected >= 4,
            f"偵測到 {len(fields)} 個欄位，命中 {detected}/{len(expected)}",
            fields_detected=len(fields),
            fields_expected=len(expected),
            issues=issues,
        )

        assert detected >= 4


# ============================================================
# 測試 4: Excel 自動回填
# ============================================================

class TestExcelAutoFill:
    """測試 auto_fill 對 Excel 的回填能力"""

    def setup_method(self):
        self.svc = get_service()

    async def _analyze_and_fill(self, content: bytes, file_name: str, values_map: dict):
        """分析 → 建立填入值 → 回填 → 讀取回填結果"""
        fields = await self.svc._deep_analyze_excel(content)

        fill_values = []
        for f in fields:
            for keyword, value in values_map.items():
                if keyword in f["field_name"]:
                    fill_values.append({"field_id": f["field_id"], "value": value})
                    break

        filled = await self.svc.auto_fill(
            file_content=content,
            file_name=file_name,
            field_map=fields,
            fill_values=fill_values,
        )

        # 讀取回填後的 Excel
        wb = load_workbook(io.BytesIO(filled))
        return fields, fill_values, filled, wb

    @pytest.mark.asyncio
    async def test_fill_simple_horizontal(self):
        """回填格式 A: 簡單水平 Excel"""
        content = create_excel_simple_horizontal()
        values = {
            "設備名稱": "馬達 A-001",
            "檢查日期": "2026-03-05",
            "設備編號": "EQ-001",
            "檢查人員": "王小明",
            "溫度": "72.5",
            "壓力": "101.3",
            "判定": "合格",
            "備註": "正常",
        }

        fields, fill_values, filled_bytes, wb = await self._analyze_and_fill(
            content, "test.xlsx", values
        )

        ws = wb.active
        # 檢查是否有值被填入
        filled_cells = 0
        for row in ws.iter_rows(min_row=1, max_row=4, max_col=4):
            for cell in row:
                if cell.value and str(cell.value) in values.values():
                    filled_cells += 1

        issues = []
        if filled_cells < len(fill_values):
            issues.append(f"填入 {filled_cells} 個，但送出 {len(fill_values)} 個值")

        report.add(
            "自動回填", "A: 簡單水平 Excel",
            filled_cells >= 4,
            f"送出 {len(fill_values)} 個值，確認填入 {filled_cells} 個",
            fill_success=filled_cells,
            fill_total=len(fill_values),
            issues=issues,
        )

        assert filled_cells >= 4, f"至少應填入 4 個值，實際 {filled_cells}"

    @pytest.mark.asyncio
    async def test_fill_multi_sheet(self):
        """回填格式 D: 多工作表 Excel"""
        content = create_excel_multi_sheet()
        values = {
            "設備名稱": "馬達 A-001",
            "設備編號": "EQ-001",
            "檢查日期": "2026-03-05",
            "溫度": "72.5",
            "判定": "合格",
        }

        fields, fill_values, filled_bytes, wb = await self._analyze_and_fill(
            content, "test.xlsx", values
        )

        # 驗證兩個 sheet 都有填入
        sheets_with_data = set()
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows(max_row=5, max_col=4):
                for cell in row:
                    if cell.value and str(cell.value) in values.values():
                        sheets_with_data.add(sheet_name)

        issues = []
        if len(sheets_with_data) < 2:
            issues.append(f"僅在 {sheets_with_data} 中填入了值")

        report.add(
            "自動回填", "D: 多工作表 Excel",
            len(sheets_with_data) >= 2,
            f"跨 {len(sheets_with_data)} 個工作表填入值",
            fill_success=len(fill_values),
            fill_total=len(fill_values),
            issues=issues,
        )

        assert len(sheets_with_data) >= 2

    @pytest.mark.asyncio
    async def test_fill_dense_labels(self):
        """回填格式 G: 密集多欄 Excel"""
        content = create_excel_dense_labels()
        values = {
            "設備名稱": "馬達 A-001",
            "型號": "AC-500",
            "規格": "三相380V",
            "檢查日期": "2026-03-05",
            "廠區": "B棟",
            "狀態": "正常",
            "溫度": "72.5",
            "壓力": "101.3",
            "電流": "15.2",
            "電壓": "380",
            "頻率": "50",
            "濕度": "45",
        }

        fields, fill_values, filled_bytes, wb = await self._analyze_and_fill(
            content, "test.xlsx", values
        )

        ws = wb.active
        filled_cells = 0
        for row in ws.iter_rows(max_row=4, max_col=6):
            for cell in row:
                if cell.value and str(cell.value) in values.values():
                    filled_cells += 1

        issues = []
        if filled_cells < 8:
            issues.append(f"密集表格僅填入 {filled_cells} 個值")

        report.add(
            "自動回填", "G: 密集多欄 Excel",
            filled_cells >= 8,
            f"送出 {len(fill_values)} 個值，確認填入 {filled_cells} 個",
            fill_success=filled_cells,
            fill_total=len(fill_values),
            issues=issues,
        )

        assert filled_cells >= 8

    @pytest.mark.asyncio
    async def test_fill_preserves_format(self):
        """驗證回填保留原始格式"""
        content = create_excel_mixed_complex()
        fields = await self.svc._deep_analyze_excel(content)

        fill_values = []
        for f in fields:
            if "檢查日期" in f["field_name"]:
                fill_values.append({"field_id": f["field_id"], "value": "2026-03-05"})

        filled = await self.svc.auto_fill(
            file_content=content,
            file_name="test.xlsx",
            field_map=fields,
            fill_values=fill_values,
        )

        wb = load_workbook(io.BytesIO(filled))
        ws = wb.active

        # 標題格式應保留
        title_font = ws['A1'].font
        is_bold = title_font.bold

        issues = []
        if not is_bold:
            issues.append("標題粗體格式遺失")

        report.add(
            "格式保留", "E: 混合複雜 Excel",
            is_bold is True,
            f"標題粗體: {is_bold}",
            issues=issues,
        )

        assert is_bold, "標題應保持粗體格式"


# ============================================================
# 測試 5: Word 自動回填
# ============================================================

class TestWordAutoFill:
    """測試 auto_fill 對 Word 的回填能力"""

    def setup_method(self):
        self.svc = get_service()

    @pytest.mark.asyncio
    async def test_fill_word_paragraph(self):
        """回填格式 H: 段落式 Word"""
        content = create_word_paragraph()
        fields = await self.svc._deep_analyze_word(content)

        values = {
            "設備名稱": "馬達 A-001",
            "設備編號": "EQ-001",
            "檢查日期": "2026-03-05",
            "檢查人員": "王小明",
            "溫度": "72.5",
            "判定": "合格",
            "備註": "正常",
        }

        fill_values = []
        for f in fields:
            for keyword, value in values.items():
                if keyword in f["field_name"]:
                    fill_values.append({"field_id": f["field_id"], "value": value})
                    break

        filled = await self.svc.auto_fill(
            file_content=content,
            file_name="test.docx",
            field_map=fields,
            fill_values=fill_values,
        )

        doc = Document(io.BytesIO(filled))
        filled_paras = 0
        for para in doc.paragraphs:
            for v in values.values():
                if v in para.text:
                    filled_paras += 1
                    break

        issues = []
        if filled_paras < 4:
            issues.append(f"僅 {filled_paras} 個段落被填入值")

        report.add(
            "自動回填", "H: 段落式 Word",
            filled_paras >= 4,
            f"送出 {len(fill_values)} 個值，確認填入 {filled_paras} 個段落",
            fill_success=filled_paras,
            fill_total=len(fill_values),
            issues=issues,
        )

        assert filled_paras >= 4

    @pytest.mark.asyncio
    async def test_fill_word_table(self):
        """回填格式 I: 表格式 Word"""
        content = create_word_table()
        fields = await self.svc._deep_analyze_word(content)

        values = {
            "設備名稱": "馬達 A-001",
            "設備編號": "EQ-001",
            "檢查日期": "2026-03-05",
            "檢查人員": "王小明",
            "位置": "B棟 2F",
            "溫度": "72.5",
            "壓力": "101.3",
            "判定": "合格",
            "備註": "正常",
        }

        fill_values = []
        for f in fields:
            for keyword, value in values.items():
                if keyword in f["field_name"]:
                    # 修補 table_index
                    if f.get("value_location") and f["value_location"].get("table_index") is None:
                        f["value_location"]["table_index"] = f["label_location"].get("table_index", 0)
                    fill_values.append({"field_id": f["field_id"], "value": value})
                    break

        filled = await self.svc.auto_fill(
            file_content=content,
            file_name="test.docx",
            field_map=fields,
            fill_values=fill_values,
        )

        doc = Document(io.BytesIO(filled))
        filled_cells = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for v in values.values():
                        if v in cell.text:
                            filled_cells += 1
                            break

        issues = []
        if filled_cells < 5:
            issues.append(f"僅 {filled_cells} 個表格儲存格被填入")

        report.add(
            "自動回填", "I: 表格式 Word",
            filled_cells >= 5,
            f"送出 {len(fill_values)} 個值，確認填入 {filled_cells} 個儲存格",
            fill_success=filled_cells,
            fill_total=len(fill_values),
            issues=issues,
        )

        assert filled_cells >= 5


# ============================================================
# 測試 6: 預覽功能
# ============================================================

class TestPreviewAutoFill:
    """測試 preview_auto_fill"""

    def setup_method(self):
        self.svc = get_service()

    @pytest.mark.asyncio
    async def test_preview_basic(self):
        """預覽基本功能"""
        field_map = [
            {"field_id": "f1", "field_name": "設備名稱", "field_type": "text",
             "value_location": {"sheet": "Sheet1", "cell": "B1"}},
            {"field_id": "f2", "field_name": "溫度", "field_type": "number",
             "value_location": {"sheet": "Sheet1", "cell": "B2"}},
            {"field_id": "f3", "field_name": "判定結果", "field_type": "checkbox",
             "value_location": None},  # 無目標位置
        ]

        fill_values = [
            {"field_id": "f1", "value": "馬達A", "confidence": 0.95, "source": "AI"},
            {"field_id": "f2", "value": "72.5", "confidence": 0.5, "source": "readings"},
            # f3 故意不提供值
        ]

        result = await self.svc.preview_auto_fill(field_map, fill_values)

        assert result["total_fields"] == 3
        assert result["filled_count"] == 2

        # 檢查 warnings
        warnings = result["warnings"]
        has_low_confidence_warning = any("信心度較低" in w for w in warnings)
        has_no_value_warning = any("無對應值" in w for w in warnings)
        has_no_target_warning = any("找不到值儲存格" in w for w in warnings)

        issues = []
        if not has_low_confidence_warning:
            issues.append("缺少低信心度警告")
        if not has_no_value_warning:
            issues.append("缺少無對應值警告")
        if not has_no_target_warning:
            issues.append("缺少無目標位置警告")

        report.add(
            "預覽功能", "預覽回填結果",
            has_low_confidence_warning and has_no_value_warning and has_no_target_warning,
            f"total={result['total_fields']}, filled={result['filled_count']}, warnings={len(warnings)}",
            issues=issues,
        )

        assert has_low_confidence_warning, "應有低信心度警告"
        assert has_no_value_warning, "應有無對應值警告"
        assert has_no_target_warning, "應有無目標位置警告"


# ============================================================
# 測試 7: 完整端到端流程
# ============================================================

class TestEndToEnd:
    """端到端流程測試: analyze → map (mock) → preview → fill"""

    def setup_method(self):
        self.svc = get_service()

    @pytest.mark.asyncio
    async def test_e2e_excel_flow(self):
        """端到端: Excel 完整流程"""
        content = create_excel_simple_horizontal()

        # Step 1: 分析結構
        result = await self.svc.analyze_structure(content, "inspection.xlsx")
        assert result["success"] is True
        field_map = result["field_map"]
        assert len(field_map) > 0

        # Step 2: 模擬 AI 映射結果 (不呼叫真實 AI)
        fill_values = []
        value_mapping = {
            "設備名稱": ("馬達 A-001", 0.95),
            "檢查日期": ("2026-03-05", 0.90),
            "設備編號": ("EQ-001", 0.85),
            "檢查人員": ("王小明", 0.80),
            "溫度": ("72.5", 0.92),
            "壓力": ("101.3", 0.88),
            "判定": ("合格", 0.95),
            "備註": ("運轉正常", 0.70),
        }
        for f in field_map:
            for keyword, (val, conf) in value_mapping.items():
                if keyword in f["field_name"]:
                    fill_values.append({
                        "field_id": f["field_id"],
                        "value": val,
                        "confidence": conf,
                        "source": f"inspection.{keyword}",
                    })
                    break

        # Step 3: 預覽
        preview = await self.svc.preview_auto_fill(field_map, fill_values)
        assert preview["total_fields"] == len(field_map)
        assert preview["filled_count"] > 0

        # Step 4: 執行回填
        filled_bytes = await self.svc.auto_fill(
            file_content=content,
            file_name="inspection.xlsx",
            field_map=field_map,
            fill_values=fill_values,
        )

        # 驗證
        wb = load_workbook(io.BytesIO(filled_bytes))
        ws = wb.active
        filled_count = 0
        for row in ws.iter_rows(max_row=4, max_col=4):
            for cell in row:
                if cell.value:
                    for _, (val, _) in value_mapping.items():
                        if str(cell.value) == val or str(cell.value) == str(float(val) if '.' in val else val):
                            filled_count += 1
                            break

        issues = []
        if filled_count < 4:
            issues.append(f"端到端流程僅確認 {filled_count} 個值")

        report.add(
            "端到端流程", "Excel 完整流程",
            filled_count >= 4,
            f"分析 {len(field_map)} 個欄位 → 映射 {len(fill_values)} 個值 → 確認填入 {filled_count} 個",
            fields_detected=len(field_map),
            fill_success=filled_count,
            fill_total=len(fill_values),
            issues=issues,
        )

        assert filled_count >= 4

    @pytest.mark.asyncio
    async def test_e2e_word_flow(self):
        """端到端: Word 完整流程"""
        content = create_word_paragraph()

        # Step 1: 分析
        result = await self.svc.analyze_structure(content, "inspection.docx")
        assert result["success"] is True
        field_map = result["field_map"]

        # Step 2: 模擬映射
        fill_values = []
        value_mapping = {
            "設備名稱": "馬達 A-001",
            "設備編號": "EQ-001",
            "檢查日期": "2026-03-05",
            "檢查人員": "王小明",
            "溫度": "72.5",
            "判定": "合格",
            "備註": "正常",
        }
        for f in field_map:
            for keyword, val in value_mapping.items():
                if keyword in f["field_name"]:
                    fill_values.append({
                        "field_id": f["field_id"],
                        "value": val,
                        "confidence": 0.9,
                        "source": "test",
                    })
                    break

        # Step 3: 預覽
        preview = await self.svc.preview_auto_fill(field_map, fill_values)

        # Step 4: 回填
        filled_bytes = await self.svc.auto_fill(
            file_content=content,
            file_name="inspection.docx",
            field_map=field_map,
            fill_values=fill_values,
        )

        doc = Document(io.BytesIO(filled_bytes))
        filled_count = 0
        for para in doc.paragraphs:
            for v in value_mapping.values():
                if v in para.text:
                    filled_count += 1
                    break

        issues = []
        if filled_count < 4:
            issues.append(f"僅確認 {filled_count} 個段落有填入值")

        report.add(
            "端到端流程", "Word 完整流程",
            filled_count >= 4,
            f"分析 {len(field_map)} 個欄位 → 映射 {len(fill_values)} 個值 → 確認填入 {filled_count} 個",
            fields_detected=len(field_map),
            fill_success=filled_count,
            fill_total=len(fill_values),
            issues=issues,
        )

        assert filled_count >= 4


# ============================================================
# 測試 8: 邊界案例
# ============================================================

class TestEdgeCases:
    """邊界案例測試"""

    def setup_method(self):
        self.svc = get_service()

    @pytest.mark.asyncio
    async def test_empty_excel(self):
        """空白 Excel"""
        wb = Workbook()
        buf = io.BytesIO()
        wb.save(buf)
        content = buf.getvalue()

        fields = await self.svc._deep_analyze_excel(content)

        report.add(
            "邊界案例", "空白 Excel",
            len(fields) == 0,
            f"偵測到 {len(fields)} 個欄位（期望 0）",
        )

        assert len(fields) == 0

    @pytest.mark.asyncio
    async def test_empty_word(self):
        """空白 Word"""
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        fields = await self.svc._deep_analyze_word(content)

        report.add(
            "邊界案例", "空白 Word",
            len(fields) == 0,
            f"偵測到 {len(fields)} 個欄位（期望 0）",
        )

        assert len(fields) == 0

    @pytest.mark.asyncio
    async def test_unsupported_file_type(self):
        """不支援的檔案格式"""
        try:
            await self.svc.analyze_structure(b"dummy", "test.pdf")
            raised = False
        except ValueError:
            raised = True

        report.add(
            "邊界案例", "不支援的格式 (PDF)",
            raised,
            "正確拋出 ValueError" if raised else "未拋出錯誤",
        )

        assert raised, "應拋出 ValueError"

    @pytest.mark.asyncio
    async def test_fill_with_no_values(self):
        """空的 fill_values"""
        content = create_excel_simple_horizontal()
        fields = await self.svc._deep_analyze_excel(content)

        filled = await self.svc.auto_fill(
            file_content=content,
            file_name="test.xlsx",
            field_map=fields,
            fill_values=[],
        )

        # 應正常回傳未修改的內容
        assert len(filled) > 0

        report.add(
            "邊界案例", "空 fill_values",
            True,
            "空值列表能正常處理，回傳原始文件",
        )

    @pytest.mark.asyncio
    async def test_fill_with_nonexistent_field_id(self):
        """fill_values 包含不存在的 field_id"""
        content = create_excel_simple_horizontal()
        fields = await self.svc._deep_analyze_excel(content)

        fill_values = [{"field_id": "nonexistent_field_xyz", "value": "test"}]

        filled = await self.svc.auto_fill(
            file_content=content,
            file_name="test.xlsx",
            field_map=fields,
            fill_values=fill_values,
        )

        assert len(filled) > 0

        report.add(
            "邊界案例", "不存在的 field_id",
            True,
            "不存在的 field_id 能正常處理，不會崩潰",
        )

    def test_convert_value_edge_cases(self):
        """值轉換邊界案例"""
        # None
        assert self.svc._convert_value(None, "number") is None
        # 空字串
        result = self.svc._convert_value("", "number")
        assert result == ""  # 無法轉為數字，返回原字串
        # 非常大的數字
        result = self.svc._convert_value("99999999.99", "number")
        assert result == 99999999.99

        report.add(
            "邊界案例", "值轉換邊界",
            True,
            "None / 空字串 / 大數字 均正確處理",
        )


# ============================================================
# 測試報告生成
# ============================================================

def generate_test_report(report_obj: TestReport) -> str:
    """產生詳細測試報告"""
    lines = []
    lines.append("=" * 80)
    lines.append("  表格自動回填功能 — 測試報告")
    lines.append(f"  測試時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 80)
    lines.append("")

    # 統計
    total = len(report_obj.results)
    passed = sum(1 for r in report_obj.results if r["passed"])
    failed = total - passed

    lines.append(f"  總測試數: {total}")
    lines.append(f"  通過: {passed}  ✓")
    lines.append(f"  失敗: {failed}  ✗")
    lines.append(f"  通過率: {passed/total*100:.1f}%" if total > 0 else "  通過率: N/A")
    lines.append("")

    # 按類別分組
    categories = {}
    for r in report_obj.results:
        cat = r["test_name"]
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(r)

    for cat, tests in categories.items():
        lines.append("-" * 70)
        lines.append(f"【{cat}】")
        lines.append("-" * 70)

        for t in tests:
            status = "✓ PASS" if t["passed"] else "✗ FAIL"
            lines.append(f"  {status}  {t['format_name']}")
            lines.append(f"         {t['details']}")

            if t["fields_detected"] or t["fields_expected"]:
                lines.append(
                    f"         欄位偵測: {t['fields_detected']}/{t['fields_expected']}"
                )
            if t["fill_success"] or t["fill_total"]:
                lines.append(
                    f"         回填成功: {t['fill_success']}/{t['fill_total']}"
                )
            if t["issues"]:
                for issue in t["issues"]:
                    lines.append(f"         ⚠ {issue}")
            lines.append("")

    # 改進建議
    lines.append("=" * 80)
    lines.append("【改進建議】")
    lines.append("=" * 80)

    all_issues = []
    for r in report_obj.results:
        if r["issues"]:
            all_issues.extend([(r["format_name"], i) for i in r["issues"]])

    if not all_issues and failed == 0:
        lines.append("  所有測試均通過，系統運作良好。")
        lines.append("")
        lines.append("  建議未來可進一步強化：")
        lines.append("  1. 支援更多檔案格式（如 PDF 表單）")
        lines.append("  2. 增強多層嵌套合併儲存格的處理")
        lines.append("  3. 支援圖片/簽名欄位的自動插入")
        lines.append("  4. 加入 AI 映射回歸測試（需要真實 API）")
        lines.append("  5. 支援更多語言的欄位關鍵字偵測（英文、日文）")
    else:
        # 針對具體問題提出建議
        issue_categories = {}
        for fmt, issue in all_issues:
            if issue not in issue_categories:
                issue_categories[issue] = []
            issue_categories[issue].append(fmt)

        idx = 1
        for issue, formats in issue_categories.items():
            lines.append(f"  {idx}. {issue}")
            lines.append(f"     影響格式: {', '.join(formats)}")
            idx += 1

        lines.append("")
        lines.append("  其他建議：")
        lines.append("  - 考慮增加對垂直表格 (below) 方向偵測的優先級")
        lines.append("  - 合併儲存格場景下值儲存格定位可進一步優化")
        lines.append("  - Word table 的 table_index 在 _find_value_cell_word 中未自動填入")
        lines.append("  - 增加對英文/日文欄位關鍵字的支援")

    lines.append("")
    lines.append("=" * 80)
    lines.append("  測試報告結束")
    lines.append("=" * 80)

    return "\n".join(lines)


# ============================================================
# pytest plugin: 測試完成後自動產生報告
# ============================================================

class ReportPlugin:
    """pytest plugin 在 session 結束後產生測試報告"""

    @staticmethod
    def pytest_sessionfinish(session, exitstatus):
        if report.results:
            report_text = generate_test_report(report)
            report_path = os.path.join(
                os.path.dirname(os.path.abspath(__file__)),
                "test_autofill_report.txt",
            )
            with open(report_path, "w", encoding="utf-8") as f:
                f.write(report_text)
            print("\n\n" + report_text)
            print(f"\n報告已儲存至: {report_path}")


# ============================================================
# 可直接執行
# ============================================================

if __name__ == "__main__":
    pytest.main(
        [__file__, "-v", "--tb=short", "-p", "no:cacheprovider"],
        plugins=[ReportPlugin()],
    )
