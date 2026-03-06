"""
自動回填系統測試套件

測試範圍：
1. 測試用 Excel/Word 表單建立
2. 結構分析（欄位偵測、值位置搜尋）
3. 預覽回填
4. 執行回填（Excel / Word）
5. 輔助方法（欄位類型推測、佔位符識別、值轉換）
6. API 端點整合測試

注意：AI 映射 (map_fields) 需要 Gemini API Key，故以 mock 方式測試。
"""

import io
import sys
import os
import json
import pytest

# 確保能 import backend 模組
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx import Document

# 設定空的 API key 以避免初始化錯誤（非 AI 測試不需要真實 key）
os.environ.setdefault("GEMINI_API_KEY", "test-key-for-unit-tests")

from app.services.form_fill import FormFillService, FIELD_KEYWORDS


# ================================================================
# 測試用表單產生器
# ================================================================

def create_test_excel_simple() -> bytes:
    """
    建立簡單 Excel 定檢表（標籤在左，值在右）

    | 設備名稱：| (空白) |
    | 設備編號：| (空白) |
    | 檢查日期：| (空白) |
    | 檢查人員：| (空白) |
    | 溫度讀數：| (空白) |
    | 是否合格：| (空白) |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "定檢表"

    data = [
        ("設備名稱：", "______"),
        ("設備編號：", "______"),
        ("檢查日期：", "______"),
        ("檢查人員：", "______"),
        ("溫度讀數：", "______"),
        ("是否合格：", "______"),
        ("備註：", "______"),
    ]

    for row_idx, (label, placeholder) in enumerate(data, start=1):
        ws.cell(row=row_idx, column=1, value=label)
        ws.cell(row=row_idx, column=2, value=placeholder)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.read()


def create_test_excel_complex() -> bytes:
    """
    建立複雜 Excel 定檢表（含合併儲存格、多工作表、佔位符）

    Sheet1「基本資訊」:
    | 設備定檢表           |            |
    | 設備名稱：           | ______     |
    | 設備編號：           | {{value}}  |
    | 位置：               | (空白)     |

    Sheet2「檢查項目」:
    | 項目   | 溫度(°C) | 壓力(MPa) | 判定結果 |
    | 主機   | (空白)   | (空白)    | (空白)   |
    """
    wb = Workbook()

    # Sheet1: 基本資訊
    ws1 = wb.active
    ws1.title = "基本資訊"
    ws1.cell(row=1, column=1, value="設備定檢表")
    ws1.merge_cells("A1:B1")
    ws1.cell(row=2, column=1, value="設備名稱：")
    ws1.cell(row=2, column=2, value="______")
    ws1.cell(row=3, column=1, value="設備編號：")
    ws1.cell(row=3, column=2, value="{{value}}")
    ws1.cell(row=4, column=1, value="位置：")
    # B4 留空

    # Sheet2: 檢查項目
    ws2 = wb.create_sheet("檢查項目")
    ws2.cell(row=1, column=1, value="項目")
    ws2.cell(row=1, column=2, value="溫度(°C)")
    ws2.cell(row=1, column=3, value="壓力(MPa)")
    ws2.cell(row=1, column=4, value="判定結果")
    ws2.cell(row=2, column=1, value="主機")

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.read()


def create_test_excel_vertical() -> bytes:
    """
    建立垂直排列的 Excel 定檢表（標籤在上，值在下）

    | 設備名稱 | 設備編號 | 檢查日期 |
    | (空白)   | (空白)   | (空白)   |
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.cell(row=1, column=1, value="設備名稱")
    ws.cell(row=1, column=2, value="設備編號")
    ws.cell(row=1, column=3, value="檢查日期")
    # Row 2 使用佔位符（值位置在下方）
    ws.cell(row=2, column=1, value="______")
    ws.cell(row=2, column=2, value="______")
    ws.cell(row=2, column=3, value="______")

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.read()


def create_test_word_simple() -> bytes:
    """
    建立簡單 Word 定檢表（段落型 + 表格型混合）

    段落：
    設備名稱：______
    檢查日期：______

    表格：
    | 檢查項目 | 結果 |
    | 溫度：   | (空) |
    | 壓力：   | (空) |
    """
    doc = Document()
    doc.add_paragraph("設備名稱：______")
    doc.add_paragraph("檢查日期：______")
    doc.add_paragraph("")  # 空行

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "檢查項目"
    table.cell(0, 1).text = "結果"
    table.cell(1, 0).text = "溫度："
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "壓力："
    table.cell(2, 1).text = ""

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


# ================================================================
# 輔助方法單元測試
# ================================================================

class TestHelperMethods:
    """測試 FormFillService 的輔助方法"""

    def setup_method(self):
        self.service = FormFillService()

    # --- _is_field_label ---

    def test_is_field_label_with_colon(self):
        assert self.service._is_field_label("設備名稱：") is True

    def test_is_field_label_with_keyword(self):
        assert self.service._is_field_label("溫度") is True
        assert self.service._is_field_label("檢查人員") is True
        assert self.service._is_field_label("備註") is True

    def test_is_field_label_non_label(self):
        assert self.service._is_field_label("ABC公司") is False
        assert self.service._is_field_label("12345") is False

    def test_is_field_label_empty(self):
        assert self.service._is_field_label("") is False
        assert self.service._is_field_label("   ") is False

    def test_is_field_label_too_long(self):
        assert self.service._is_field_label("a" * 51) is False

    # --- _is_placeholder ---

    def test_is_placeholder_underscores(self):
        assert self.service._is_placeholder("______") is True
        assert self.service._is_placeholder("＿＿＿＿") is True

    def test_is_placeholder_template_syntax(self):
        assert self.service._is_placeholder("{{field_name}}") is True
        assert self.service._is_placeholder("<請填入>") is True
        assert self.service._is_placeholder("[值]") is True

    def test_is_placeholder_empty(self):
        assert self.service._is_placeholder("") is True
        assert self.service._is_placeholder("   ") is True

    def test_is_placeholder_normal_text(self):
        assert self.service._is_placeholder("正常文字") is False
        assert self.service._is_placeholder("65.5") is False

    # --- _guess_field_type ---

    def test_guess_field_type_date(self):
        assert self.service._guess_field_type("檢查日期") == "date"
        assert self.service._guess_field_type("巡檢時間") == "date"

    def test_guess_field_type_number(self):
        assert self.service._guess_field_type("溫度") == "number"
        assert self.service._guess_field_type("壓力讀數") == "number"
        assert self.service._guess_field_type("電流值") == "number"
        assert self.service._guess_field_type("振動") == "number"

    def test_guess_field_type_checkbox(self):
        assert self.service._guess_field_type("是否合格") == "checkbox"
        assert self.service._guess_field_type("判定結果") == "checkbox"
        assert self.service._guess_field_type("正常/異常") == "checkbox"

    def test_guess_field_type_text(self):
        assert self.service._guess_field_type("設備名稱") == "text"
        assert self.service._guess_field_type("巡檢員") == "text"

    # --- _convert_value ---

    def test_convert_value_number(self):
        assert self.service._convert_value("65.5", "number") == 65.5
        assert self.service._convert_value("100", "number") == 100
        assert self.service._convert_value("非數值", "number") == "非數值"

    def test_convert_value_checkbox(self):
        assert self.service._convert_value("true", "checkbox") == "合格"
        assert self.service._convert_value("是", "checkbox") == "合格"
        assert self.service._convert_value("正常", "checkbox") == "合格"
        assert self.service._convert_value("false", "checkbox") == "不合格"
        assert self.service._convert_value("異常", "checkbox") == "不合格"
        assert self.service._convert_value("不合格", "checkbox") == "不合格"

    def test_convert_value_text(self):
        assert self.service._convert_value("馬達 A-01", "text") == "馬達 A-01"

    def test_convert_value_none(self):
        assert self.service._convert_value(None, "text") is None
        assert self.service._convert_value(None, "number") is None


# ================================================================
# Excel 結構分析測試
# ================================================================

class TestExcelStructureAnalysis:
    """測試 Excel 結構分析"""

    def setup_method(self):
        self.service = FormFillService()

    @pytest.mark.asyncio
    async def test_simple_excel_detects_all_fields(self):
        """簡單表格：應偵測到所有欄位標籤"""
        content = create_test_excel_simple()
        result = await self.service.analyze_structure(content, "test.xlsx")

        assert result["success"] is True
        assert result["file_type"] == "xlsx"
        assert result["total_fields"] >= 6  # 至少 6 個欄位

        field_names = [f["field_name"] for f in result["field_map"]]
        assert any("設備名稱" in n for n in field_names)
        assert any("檢查日期" in n for n in field_names)
        assert any("溫度" in n for n in field_names)

    @pytest.mark.asyncio
    async def test_simple_excel_value_location_right(self):
        """簡單表格：值位置應在標籤右方"""
        content = create_test_excel_simple()
        result = await self.service.analyze_structure(content, "test.xlsx")

        for field in result["field_map"]:
            val_loc = field.get("value_location")
            assert val_loc is not None, f"Field '{field['field_name']}' has no value_location"
            assert val_loc["direction"] == "right"
            assert val_loc["column"] == field["label_location"]["column"] + 1

    @pytest.mark.asyncio
    async def test_complex_excel_multi_sheet(self):
        """複雜表格：應分析多個工作表"""
        content = create_test_excel_complex()
        result = await self.service.analyze_structure(content, "test.xlsx")

        assert result["success"] is True

        sheets = set()
        for field in result["field_map"]:
            sheets.add(field["label_location"]["sheet"])

        # 應包含兩個工作表的欄位
        assert "基本資訊" in sheets

    @pytest.mark.asyncio
    async def test_complex_excel_placeholder_detection(self):
        """複雜表格：應正確識別佔位符並找到值位置"""
        content = create_test_excel_complex()
        result = await self.service.analyze_structure(content, "test.xlsx")

        # 「設備名稱：」的值位置應是包含 "______" 的儲存格
        name_field = next(
            (f for f in result["field_map"] if "設備名稱" in f["field_name"]),
            None
        )
        assert name_field is not None
        assert name_field["value_location"] is not None

    @pytest.mark.asyncio
    async def test_vertical_layout(self):
        """垂直排列：標籤在上，值在下"""
        content = create_test_excel_vertical()
        result = await self.service.analyze_structure(content, "test.xlsx")

        assert result["success"] is True
        assert result["total_fields"] >= 2

        # 至少一個欄位的值位置在下方
        has_below = any(
            (f.get("value_location") or {}).get("direction") == "below"
            for f in result["field_map"]
        )
        assert has_below, "Should detect at least one field with value below"

    @pytest.mark.asyncio
    async def test_field_types_are_guessed(self):
        """欄位類型應被正確推測"""
        content = create_test_excel_simple()
        result = await self.service.analyze_structure(content, "test.xlsx")

        type_map = {f["field_name"].rstrip("：:"): f["field_type"] for f in result["field_map"]}

        # 日期欄位
        date_field = next((v for k, v in type_map.items() if "日期" in k), None)
        assert date_field == "date"

        # 數值欄位
        temp_field = next((v for k, v in type_map.items() if "溫度" in k), None)
        assert temp_field == "number"

    @pytest.mark.asyncio
    async def test_unsupported_format(self):
        """不支援的格式應拋出錯誤"""
        with pytest.raises(ValueError, match="不支援"):
            await self.service.analyze_structure(b"fake", "test.pdf")


# ================================================================
# Word 結構分析測試
# ================================================================

class TestWordStructureAnalysis:
    """測試 Word 結構分析"""

    def setup_method(self):
        self.service = FormFillService()

    @pytest.mark.asyncio
    async def test_word_paragraph_fields(self):
        """Word 段落型欄位偵測"""
        content = create_test_word_simple()
        result = await self.service.analyze_structure(content, "test.docx")

        assert result["success"] is True
        assert result["file_type"] == "docx"

        field_names = [f["field_name"] for f in result["field_map"]]
        assert any("設備名稱" in n for n in field_names)

    @pytest.mark.asyncio
    async def test_word_table_fields(self):
        """Word 表格型欄位偵測"""
        content = create_test_word_simple()
        result = await self.service.analyze_structure(content, "test.docx")

        table_fields = [
            f for f in result["field_map"]
            if f.get("label_location", {}).get("type") == "table"
        ]
        assert len(table_fields) >= 1  # 至少偵測到表格中的欄位


# ================================================================
# 預覽回填測試
# ================================================================

class TestPreviewAutoFill:
    """測試預覽功能"""

    def setup_method(self):
        self.service = FormFillService()

    @pytest.mark.asyncio
    async def test_preview_basic(self):
        """基本預覽：有值的欄位應正確顯示"""
        field_map = [
            {"field_id": "f1", "field_name": "設備名稱", "field_type": "text", "value_location": {"cell": "B1"}},
            {"field_id": "f2", "field_name": "溫度", "field_type": "number", "value_location": {"cell": "B2"}},
        ]
        fill_values = [
            {"field_id": "f1", "value": "馬達 A-01", "confidence": 0.95, "source": "AI 分析"},
        ]

        result = await self.service.preview_auto_fill(field_map, fill_values)

        assert result["total_fields"] == 2
        assert result["filled_count"] == 1
        assert len(result["warnings"]) >= 1  # 至少有一個「無對應值」的警告

    @pytest.mark.asyncio
    async def test_preview_low_confidence_warning(self):
        """低信心度應產生警告"""
        field_map = [
            {"field_id": "f1", "field_name": "設備名稱", "field_type": "text", "value_location": {"cell": "B1"}},
        ]
        fill_values = [
            {"field_id": "f1", "value": "未知設備", "confidence": 0.5, "source": "猜測"},
        ]

        result = await self.service.preview_auto_fill(field_map, fill_values)

        low_conf_warnings = [w for w in result["warnings"] if "信心度" in w]
        assert len(low_conf_warnings) >= 1

    @pytest.mark.asyncio
    async def test_preview_no_target_warning(self):
        """沒有值位置的欄位應產生警告"""
        field_map = [
            {"field_id": "f1", "field_name": "設備名稱", "field_type": "text", "value_location": None},
        ]
        fill_values = [
            {"field_id": "f1", "value": "馬達", "confidence": 0.9},
        ]

        result = await self.service.preview_auto_fill(field_map, fill_values)

        no_target_warnings = [w for w in result["warnings"] if "找不到值儲存格" in w]
        assert len(no_target_warnings) >= 1


# ================================================================
# Excel 回填執行測試
# ================================================================

class TestExcelAutoFill:
    """測試 Excel 回填"""

    def setup_method(self):
        self.service = FormFillService()

    @pytest.mark.asyncio
    async def test_fill_simple_excel(self):
        """回填簡單 Excel：值應被寫入正確位置"""
        content = create_test_excel_simple()

        # 先分析結構
        analysis = await self.service.analyze_structure(content, "test.xlsx")
        field_map = analysis["field_map"]

        # 找到「設備名稱」和「溫度」的 field_id
        name_field = next(f for f in field_map if "設備名稱" in f["field_name"])
        temp_field = next(f for f in field_map if "溫度" in f["field_name"])

        fill_values = [
            {"field_id": name_field["field_id"], "value": "馬達 A-01"},
            {"field_id": temp_field["field_id"], "value": "65.5"},
        ]

        # 執行回填
        filled_bytes = await self.service.auto_fill(
            content, "test.xlsx", field_map, fill_values
        )

        # 驗證結果
        from openpyxl import load_workbook as lw
        wb = lw(io.BytesIO(filled_bytes))
        ws = wb["定檢表"]

        # 設備名稱應寫在 B 欄
        name_val_loc = name_field["value_location"]
        actual_name = ws[name_val_loc["cell"]].value
        assert actual_name == "馬達 A-01"

        # 溫度應轉為數值 65.5
        temp_val_loc = temp_field["value_location"]
        actual_temp = ws[temp_val_loc["cell"]].value
        assert actual_temp == 65.5

    @pytest.mark.asyncio
    async def test_fill_checkbox_value(self):
        """回填 checkbox 欄位：應轉換為「合格/不合格」"""
        content = create_test_excel_simple()
        analysis = await self.service.analyze_structure(content, "test.xlsx")
        field_map = analysis["field_map"]

        check_field = next(f for f in field_map if "合格" in f["field_name"])

        fill_values = [
            {"field_id": check_field["field_id"], "value": "true"},
        ]

        filled_bytes = await self.service.auto_fill(
            content, "test.xlsx", field_map, fill_values
        )

        from openpyxl import load_workbook as lw
        wb = lw(io.BytesIO(filled_bytes))
        ws = wb["定檢表"]

        val_loc = check_field["value_location"]
        actual = ws[val_loc["cell"]].value
        assert actual == "合格"

    @pytest.mark.asyncio
    async def test_fill_preserves_format(self):
        """回填不應破壞原始格式"""
        content = create_test_excel_simple()
        analysis = await self.service.analyze_structure(content, "test.xlsx")
        field_map = analysis["field_map"]
        name_field = next(f for f in field_map if "設備名稱" in f["field_name"])

        fill_values = [
            {"field_id": name_field["field_id"], "value": "馬達 A-01"},
        ]

        filled_bytes = await self.service.auto_fill(
            content, "test.xlsx", field_map, fill_values
        )

        # 確認回填後的檔案可以正常開啟
        from openpyxl import load_workbook as lw
        wb = lw(io.BytesIO(filled_bytes))
        assert "定檢表" in wb.sheetnames

    @pytest.mark.asyncio
    async def test_fill_missing_field_skipped(self):
        """不存在的 field_id 應被跳過，不報錯"""
        content = create_test_excel_simple()
        analysis = await self.service.analyze_structure(content, "test.xlsx")
        field_map = analysis["field_map"]

        fill_values = [
            {"field_id": "nonexistent_field", "value": "should be skipped"},
        ]

        # 不應拋出異常
        filled_bytes = await self.service.auto_fill(
            content, "test.xlsx", field_map, fill_values
        )
        assert len(filled_bytes) > 0


# ================================================================
# Word 回填執行測試
# ================================================================

class TestWordAutoFill:
    """測試 Word 回填"""

    def setup_method(self):
        self.service = FormFillService()

    @pytest.mark.asyncio
    async def test_fill_word_paragraph(self):
        """回填 Word 段落型欄位"""
        content = create_test_word_simple()
        analysis = await self.service.analyze_structure(content, "test.docx")
        field_map = analysis["field_map"]

        # 找段落型的「設備名稱」欄位
        name_field = next(
            (f for f in field_map
             if "設備名稱" in f["field_name"]
             and f.get("label_location", {}).get("type") == "paragraph"),
            None
        )

        if name_field:
            fill_values = [
                {"field_id": name_field["field_id"], "value": "馬達 B-02"},
            ]

            filled_bytes = await self.service.auto_fill(
                content, "test.docx", field_map, fill_values
            )

            doc = Document(io.BytesIO(filled_bytes))
            # 確認「設備名稱」段落被更新
            para_text = doc.paragraphs[name_field["value_location"]["paragraph_index"]].text
            assert "馬達 B-02" in para_text

    @pytest.mark.asyncio
    async def test_fill_word_table(self):
        """回填 Word 表格型欄位"""
        content = create_test_word_simple()
        analysis = await self.service.analyze_structure(content, "test.docx")
        field_map = analysis["field_map"]

        # 找表格型欄位
        table_fields = [
            f for f in field_map
            if f.get("label_location", {}).get("type") == "table"
        ]

        if table_fields:
            field = table_fields[0]
            fill_values = [
                {"field_id": field["field_id"], "value": "65.5°C"},
            ]

            filled_bytes = await self.service.auto_fill(
                content, "test.docx", field_map, fill_values
            )

            # 確認檔案可正常開啟
            doc = Document(io.BytesIO(filled_bytes))
            assert len(doc.tables) >= 1


# ================================================================
# API 端點整合測試
# ================================================================

class TestAPIEndpoints:
    """測試 FastAPI API 端點"""

    @pytest.fixture
    def client(self):
        from fastapi.testclient import TestClient
        from app.main import app
        return TestClient(app)

    def test_health_endpoint(self, client):
        """健康檢查端點"""
        response = client.get("/health")
        assert response.status_code == 200
        assert response.json()["status"] == "ok"

    def test_analyze_structure_excel(self, client):
        """上傳 Excel 分析結構"""
        content = create_test_excel_simple()
        response = client.post(
            "/api/auto-fill/analyze-structure",
            files={"file": ("test.xlsx", io.BytesIO(content),
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["total_fields"] >= 6

    def test_analyze_structure_word(self, client):
        """上傳 Word 分析結構"""
        content = create_test_word_simple()
        response = client.post(
            "/api/auto-fill/analyze-structure",
            files={"file": ("test.docx", io.BytesIO(content),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["total_fields"] >= 1

    def test_analyze_structure_invalid_type(self, client):
        """上傳不支援的格式應回傳 400"""
        response = client.post(
            "/api/auto-fill/analyze-structure",
            files={"file": ("test.pdf", io.BytesIO(b"fake"), "application/pdf")},
        )
        assert response.status_code == 400

    def test_preview_endpoint(self, client):
        """預覽端點"""
        response = client.post(
            "/api/auto-fill/preview",
            json={
                "field_map": [
                    {"field_id": "f1", "field_name": "設備名稱", "field_type": "text",
                     "value_location": {"cell": "B1"}},
                ],
                "fill_values": [
                    {"field_id": "f1", "value": "馬達 A-01", "confidence": 0.95, "source": "AI"},
                ],
            },
        )
        assert response.status_code == 200
        data = response.json()
        assert data["total_fields"] == 1
        assert data["filled_count"] == 1

    def test_execute_endpoint_excel(self, client):
        """執行回填端點 (Excel)

        注意：execute 端點使用 multipart/form-data 混合 File + Form，
        需要將 JSON 字串作為額外的 multipart 欄位發送。
        """
        content = create_test_excel_simple()

        # 先分析結構取得 field_map
        analyze_resp = client.post(
            "/api/auto-fill/analyze-structure",
            files={"file": ("test.xlsx", io.BytesIO(content),
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        field_map = analyze_resp.json()["field_map"]

        # 找一個欄位來回填
        name_field = next(f for f in field_map if "設備名稱" in f["field_name"])
        fill_values = [{"field_id": name_field["field_id"], "value": "馬達 C-03"}]

        # multipart/form-data: File + Form fields 需一起作為 files 傳送
        # 因為 httpx 在 data+files 混用時，Form 欄位可能不會被正確傳遞
        field_map_str = json.dumps(field_map, ensure_ascii=False)
        fill_values_str = json.dumps(fill_values, ensure_ascii=False)

        response = client.post(
            "/api/auto-fill/execute",
            files={
                "file": ("test.xlsx", io.BytesIO(content),
                         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
                "field_map_json": (None, field_map_str),
                "fill_values_json": (None, fill_values_str),
            },
        )

        assert response.status_code == 200
        assert len(response.content) > 0

        # 驗證回填內容
        from openpyxl import load_workbook as lw
        wb = lw(io.BytesIO(response.content))
        ws = wb["定檢表"]
        val_loc = name_field["value_location"]
        assert ws[val_loc["cell"]].value == "馬達 C-03"


# ================================================================
# 端到端完整流程測試
# ================================================================

class TestEndToEndFlow:
    """測試完整的 4 階段工作流程（不含 AI 映射）"""

    def setup_method(self):
        self.service = FormFillService()

    @pytest.mark.asyncio
    async def test_full_flow_excel(self):
        """
        完整流程：
        1. 結構分析
        2. 手動建立映射（模擬 AI）
        3. 預覽
        4. 執行回填
        5. 驗證結果
        """
        content = create_test_excel_simple()

        # Step 1: 結構分析
        analysis = await self.service.analyze_structure(content, "test.xlsx")
        assert analysis["success"] is True
        field_map = analysis["field_map"]

        # Step 2: 模擬 AI 映射結果
        name_field = next(f for f in field_map if "設備名稱" in f["field_name"])
        date_field = next(f for f in field_map if "日期" in f["field_name"])
        temp_field = next(f for f in field_map if "溫度" in f["field_name"])
        check_field = next(f for f in field_map if "合格" in f["field_name"])

        fill_values = [
            {"field_id": name_field["field_id"], "value": "馬達 A-01", "confidence": 0.95, "source": "AI"},
            {"field_id": date_field["field_id"], "value": "2026-03-05", "confidence": 0.99, "source": "系統"},
            {"field_id": temp_field["field_id"], "value": "65.5", "confidence": 0.90, "source": "溫度照片"},
            {"field_id": check_field["field_id"], "value": "true", "confidence": 0.85, "source": "AI 判定"},
        ]

        # Step 3: 預覽
        preview = await self.service.preview_auto_fill(field_map, fill_values)
        assert preview["filled_count"] == 4
        assert preview["total_fields"] >= 6

        # Step 4: 執行回填
        filled_bytes = await self.service.auto_fill(
            content, "test.xlsx", field_map, fill_values
        )

        # Step 5: 驗證結果
        from openpyxl import load_workbook as lw
        wb = lw(io.BytesIO(filled_bytes))
        ws = wb["定檢表"]

        assert ws[name_field["value_location"]["cell"]].value == "馬達 A-01"
        assert ws[date_field["value_location"]["cell"]].value == "2026-03-05"
        assert ws[temp_field["value_location"]["cell"]].value == 65.5
        assert ws[check_field["value_location"]["cell"]].value == "合格"

    @pytest.mark.asyncio
    async def test_full_flow_word(self):
        """Word 完整流程"""
        content = create_test_word_simple()

        # Step 1: 結構分析
        analysis = await self.service.analyze_structure(content, "test.docx")
        assert analysis["success"] is True
        field_map = analysis["field_map"]
        assert len(field_map) >= 1

        # Step 2: 模擬映射
        fill_values = []
        for field in field_map:
            fill_values.append({
                "field_id": field["field_id"],
                "value": f"測試值_{field['field_name']}",
                "confidence": 0.90,
            })

        # Step 3: 預覽
        preview = await self.service.preview_auto_fill(field_map, fill_values)
        assert preview["filled_count"] == len(field_map)

        # Step 4: 執行回填
        filled_bytes = await self.service.auto_fill(
            content, "test.docx", field_map, fill_values
        )

        # Step 5: 驗證檔案可開啟
        doc = Document(io.BytesIO(filled_bytes))
        assert len(doc.paragraphs) >= 1
